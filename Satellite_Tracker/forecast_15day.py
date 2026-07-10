"""
15-day Pakistan-airspace forecast for high-resolution EO satellites.

Generates:
  - forecast_15day.json   (canonical artefact consumed by dashboard + API)
  - documents/Pakistan_15Day_Forecast_<YYYY-MM-DD>.docx  (printable report)

For each of the next 15 days, every curated <=1 m EO satellite is propagated
at 60-second steps across the Pakistan box, every contiguous over-Pakistan
arc is recorded as a "crossing", per-day observation windows are merged from
those crossings, and the complement (gap) becomes the day's "blind windows".

Indian (ISRO) satellites are flagged and rendered in red.
"""

from __future__ import annotations
import datetime
import json
import math
import os
import sys
import time
from datetime import timezone, timedelta
from typing import Iterable

# Make in-folder imports work whether you run this file directly or import it.
_HERE = os.path.dirname(os.path.abspath(__file__))
if _HERE not in sys.path:
    sys.path.insert(0, _HERE)

from hires_eo_satellites import (
    HIRES_EO_SATELLITES, INDIAN_NORAD_IDS, TILT_SPECS,
    altitude_km as _alt_km, max_tilt_deg as _tilt_deg, standoff_km as _so_km,
)
from celestrak_tle import fetch_hires_eo_tles
from pakistan_strategic_sites import (
    STRATEGIC_SITES, TIER1_SITES, PRIORITY_CITIES, find_targeted_sites,
)

# ── Constants ────────────────────────────────────────────────────────────────
PAK_LAT_MIN, PAK_LAT_MAX = 23.5, 37.5
PAK_LON_MIN, PAK_LON_MAX = 60.5, 77.5

# Per-satellite tilt range REPLACES the old single ~300 km global buffer
# (FEAT-003). Each satellite's standoff is its altitude * tan(max_tilt) —
# data lives in hires_eo_satellites.TILT_SPECS. For any sat without a spec
# we fall back to the legacy 300 km for safety.
LEGACY_TILT_KM = 300.0
TILT_BUFFER_LAT_DEG = LEGACY_TILT_KM / 111.0   # kept for back-compat callers
TILT_BUFFER_LON_DEG = LEGACY_TILT_KM /  96.0   # kept for back-compat callers
# 1 degree of latitude is ~111 km. 1 degree of longitude at ~30N is ~96 km.
KM_PER_DEG_LAT = 111.0
KM_PER_DEG_LON_30N = 96.0


def _per_sat_buffer_deg(norad: str | int) -> tuple[float, float]:
    """Convert this satellite's standoff_km into a (lat_deg, lon_deg) buffer."""
    so = _so_km(norad)
    if so is None:
        so = LEGACY_TILT_KM
    return (so / KM_PER_DEG_LAT, so / KM_PER_DEG_LON_30N)

FORECAST_DAYS = 15
STEP_SECONDS = 60

BASE_DIR = _HERE
FORECAST_JSON = os.path.join(BASE_DIR, "forecast_15day.json")
DOCS_DIR = os.path.join(BASE_DIR, "documents")
os.makedirs(DOCS_DIR, exist_ok=True)


def _cleanup_pngs(paths):
    """Delete intermediate coverage/map PNGs once they're embedded in a PDF.
    The operator doesn't want these images left on disk — they bloated
    documents/ and every backup snapshot (~1 GB across 30 snapshots)."""
    for _p in (paths or ()):
        try:
            os.remove(_p)
        except OSError:
            pass


# ── SGP4 propagation — prefer the real library, fall back to the hand-rolled one
# ────────────────────────────────────────────────────────────────────────────
try:
    from sgp4.api import Satrec, jday
    _HAVE_SGP4 = True
except ImportError:  # pragma: no cover — fallback path
    _HAVE_SGP4 = False


def _propagate_sgp4(line1: str, line2: str, dt: datetime.datetime
                    ) -> tuple[float, float, float] | None:
    """Return (lat_deg, lon_deg, alt_km) using the official sgp4 library."""
    try:
        sat = Satrec.twoline2rv(line1, line2)
        jd, fr = jday(dt.year, dt.month, dt.day,
                      dt.hour, dt.minute, dt.second + dt.microsecond / 1e6)
        e, r, _v = sat.sgp4(jd, fr)
        if e != 0:
            return None
        x, y, z = r  # TEME km

        # GMST → ECEF
        j2k = datetime.datetime(2000, 1, 1, 12, 0, 0, tzinfo=timezone.utc)
        d = (dt - j2k).total_seconds() / 86400.0
        gmst = math.radians((280.46061837 + 360.98564736629 * d) % 360)
        cg, sg = math.cos(gmst), math.sin(gmst)
        xe = x * cg + y * sg
        ye = -x * sg + y * cg
        ze = z
        lon = math.degrees(math.atan2(ye, xe))
        lat = math.degrees(math.atan2(ze, math.sqrt(xe ** 2 + ye ** 2)))
        alt = math.sqrt(x ** 2 + y ** 2 + z ** 2) - 6371.0
        return lat, lon, alt
    except (ValueError, RuntimeError, OverflowError):
        return None


def _propagate_handrolled(line1: str, line2: str, dt: datetime.datetime
                          ) -> tuple[float, float, float] | None:
    """Backup propagator — Kepler-only, less accurate but no dependency."""
    try:
        ey = int(line1[18:20]); ed = float(line1[20:32])
        yr = ey + (2000 if ey < 57 else 1900)
        epoch = datetime.datetime(yr, 1, 1, tzinfo=timezone.utc) + timedelta(days=ed - 1)
        inc  = math.radians(float(line2[8:16]))
        raan = math.radians(float(line2[17:25]))
        ecc  = float("0." + line2[26:33].strip())
        argp = math.radians(float(line2[34:42]))
        ma0  = math.radians(float(line2[43:51]))
        n    = float(line2[52:63])             # rev / day
        n_rad_min = n * 2 * math.pi / 1440.0
        n_rad_sec = n * 2 * math.pi / 86400.0
        a = (398600.4418 / (n_rad_sec ** 2)) ** (1 / 3)

        dt_min = (dt - epoch).total_seconds() / 60.0
        M = (ma0 + n_rad_min * dt_min) % (2 * math.pi)

        E = M
        for _ in range(10):
            E = E - (E - ecc * math.sin(E) - M) / (1 - ecc * math.cos(E))

        nu = 2 * math.atan2(math.sqrt(1 + ecc) * math.sin(E / 2),
                            math.sqrt(1 - ecc) * math.cos(E / 2))
        r = a * (1 - ecc * math.cos(E))
        alt = r - 6371.0

        xo = r * math.cos(nu); yo = r * math.sin(nu)
        cr, sr = math.cos(raan), math.sin(raan)
        ci, si = math.cos(inc), math.sin(inc)
        cw, sw = math.cos(argp), math.sin(argp)
        x = (cr * cw - sr * sw * ci) * xo + (-cr * sw - sr * cw * ci) * yo
        y = (sr * cw + cr * sw * ci) * xo + (-sr * sw + cr * cw * ci) * yo
        z = (sw * si) * xo + (cw * si) * yo

        j2k = datetime.datetime(2000, 1, 1, 12, 0, 0, tzinfo=timezone.utc)
        d = (dt - j2k).total_seconds() / 86400.0
        gmst = math.radians((280.46061837 + 360.98564736629 * d) % 360)
        xe = x * math.cos(gmst) + y * math.sin(gmst)
        ye = -x * math.sin(gmst) + y * math.cos(gmst)
        ze = z
        lon = math.degrees(math.atan2(ye, xe))
        lat = math.degrees(math.atan2(ze, math.sqrt(xe ** 2 + ye ** 2)))
        return lat, lon, alt
    except (ValueError, ZeroDivisionError, OverflowError):
        return None


def propagate(line1: str, line2: str, dt: datetime.datetime
              ) -> tuple[float, float, float] | None:
    if _HAVE_SGP4:
        result = _propagate_sgp4(line1, line2, dt)
        if result is not None:
            return result
    return _propagate_handrolled(line1, line2, dt)


def _over_pak(lat: float, lon: float) -> bool:
    """Sub-point is directly over Pakistan."""
    return (PAK_LAT_MIN <= lat <= PAK_LAT_MAX
            and PAK_LON_MIN <= lon <= PAK_LON_MAX)


def _in_tilt_zone(lat: float, lon: float,
                  buf_lat_deg: float = TILT_BUFFER_LAT_DEG,
                  buf_lon_deg: float = TILT_BUFFER_LON_DEG) -> bool:
    """
    Sub-point is within this satellite's tilt-standoff reach of the
    Pakistan box (or directly over). The buffer is per-satellite —
    callers should pass `_per_sat_buffer_deg(norad)`; the default global
    values are kept only for back-compat with code that doesn't yet know
    its NORAD.
    """
    return (PAK_LAT_MIN - buf_lat_deg <= lat <= PAK_LAT_MAX + buf_lat_deg
            and PAK_LON_MIN - buf_lon_deg <= lon <= PAK_LON_MAX + buf_lon_deg)


# ── Solar elevation — used to gate optical-EO crossings to daylight only ───
def _solar_elevation_deg(lat: float, lon: float,
                         dt: datetime.datetime) -> float:
    """
    Approximate solar elevation angle (degrees above horizon) at a given
    geographic point and UTC time. Standard NOAA solar-position formulas,
    accurate to ~0.1 degrees — more than enough to decide day vs night.
    """
    # Day of year (fractional)
    jd = (dt - datetime.datetime(dt.year, 1, 1, tzinfo=timezone.utc)).total_seconds() / 86400.0 + 1
    # Solar declination
    g = math.radians(360.0 / 365.25 * (jd - 81))
    decl = math.radians(23.44) * math.sin(g)
    # Equation of time (minutes)
    eot = 9.87 * math.sin(2 * g) - 7.53 * math.cos(g) - 1.5 * math.sin(g)
    # Solar time (hours), local
    utc_hours = dt.hour + dt.minute / 60 + dt.second / 3600
    solar_time = utc_hours + lon / 15.0 + eot / 60.0
    # Hour angle (degrees from solar noon)
    ha = math.radians(15.0 * (solar_time - 12.0))
    phi = math.radians(lat)
    sin_alt = (math.sin(phi) * math.sin(decl)
               + math.cos(phi) * math.cos(decl) * math.cos(ha))
    sin_alt = max(-1.0, min(1.0, sin_alt))
    return math.degrees(math.asin(sin_alt))


# Minimum sun elevation for usable optical imagery. 0 deg = just-above
# horizon (we accept civil-twilight passes too — sat is over a sunlit
# Earth from its altitude). Use 5 deg if you want stricter "real daylight".
DAYLIGHT_MIN_ELEV_DEG = 0.0


def _is_optical(sensor: str) -> bool:
    """True if the sensor needs daylight (i.e., not SAR)."""
    return "SAR" not in sensor.upper()


# ── Crossing detection (per satellite per day) ──────────────────────────────
def _day_crossings(tle: dict, day: datetime.date) -> list[dict]:
    """
    Detect all near-Pakistan arcs for a single satellite in one UTC day.

    An "arc" is a contiguous run of timesteps where the satellite's sub-point
    is inside the Pakistan box OR within the 300 km tilt buffer around it.

    Each arc is classified:
      pass_type = "overhead"   — sub-point entered Pakistan box at any moment
      pass_type = "tilt-range" — never entered the box, only the 300 km buffer
                                  (optical sensor could still tilt + image)
    Daylight filter: OPTICAL arcs whose midpoint has the sun below the horizon
    are discarded (they can't image in the dark). SAR arcs are kept day AND
    night, because radar images regardless of sunlight.
    """
    line1, line2 = tle["line1"], tle["line2"]
    # 'day' is a Pakistan calendar date. A PKT day runs 00:00–24:00 PKT, which
    # in UTC is 19:00 (previous day) → 19:00. Shift the window back 5h so each
    # bucket is a true Pakistan day, and pass times (rendered as UTC+5) fall on
    # the day the user actually sees them.
    start = (datetime.datetime(day.year, day.month, day.day, tzinfo=timezone.utc)
             - timedelta(hours=5))
    end   = start + timedelta(days=1)
    step  = timedelta(seconds=STEP_SECONDS)

    # Per-satellite tilt buffer (replaces the legacy global 300 km).
    sat_buf_lat, sat_buf_lon = _per_sat_buffer_deg(tle["norad_id"])

    crossings: list[dict] = []
    in_arc = False
    entry_t: datetime.datetime | None = None
    entry_pos: tuple[float, float] | None = None
    arc_positions: list[tuple[datetime.datetime, float, float, float]] = []
    arc_min_dist_deg = 1e9  # min distance to Pakistan box edge for the arc
    arc_was_overhead = False

    def _close(end_time: datetime.datetime):
        nonlocal in_arc, arc_positions, arc_was_overhead, arc_min_dist_deg
        if entry_t is None or entry_pos is None or not arc_positions:
            return
        pass_type = "overhead" if arc_was_overhead else "tilt-range"
        crossings.append(_pack_crossing(
            tle, entry_t, end_time, entry_pos, arc_positions,
            pass_type, arc_min_dist_deg))
        in_arc = False
        arc_positions = []
        arc_was_overhead = False
        arc_min_dist_deg = 1e9

    t = start
    while t <= end:
        pos = propagate(line1, line2, t)
        if pos is None:
            if in_arc:
                _close(t)
            t += step
            continue

        lat, lon, alt = pos
        in_tilt = _in_tilt_zone(lat, lon, sat_buf_lat, sat_buf_lon)
        over    = _over_pak(lat, lon) if in_tilt else False

        if in_tilt:
            if not in_arc:
                in_arc = True
                entry_t = t
                entry_pos = (lat, lon)
                arc_positions = []
                arc_was_overhead = False
                arc_min_dist_deg = 1e9
            if over:
                arc_was_overhead = True
            # track closest approach to the box (0 = inside)
            d = _dist_to_pak_box_deg(lat, lon)
            if d < arc_min_dist_deg:
                arc_min_dist_deg = d
            arc_positions.append((t, lat, lon, alt))
        elif in_arc:
            _close(t)
        t += step

    if in_arc:
        _close(end)

    # Daylight filter applies to OPTICAL sensors only. Optical needs sunlight
    # to image, so its night passes are not useful coverage and are dropped.
    # SAR is an all-weather day/night radar: a SAR satellite over Pakistan at
    # 2 a.m. is still imaging, so SAR passes are kept regardless of sun
    # elevation. (This filter used to be universal, which wrongly deleted every
    # night pass — SAR included — and made the small hours look permanently
    # blind with a bogus ~8 h "longest gap".)
    def _keep_crossing(c: dict) -> bool:
        if "SAR" in (c.get("sensor") or "").upper():
            return True                      # radar sees day and night
        return c.get("sun_elev_deg", 0) > DAYLIGHT_MIN_ELEV_DEG
    crossings = [c for c in crossings if _keep_crossing(c)]
    return crossings


def _dist_to_pak_box_deg(lat: float, lon: float) -> float:
    """Approx degrees from (lat,lon) to nearest point on the Pakistan box.
    0 if inside. Used only as a sort key — no need for great-circle math."""
    dlat = max(PAK_LAT_MIN - lat, 0.0, lat - PAK_LAT_MAX)
    dlon = max(PAK_LON_MIN - lon, 0.0, lon - PAK_LON_MAX)
    return math.sqrt(dlat * dlat + dlon * dlon)


def _pack_crossing(tle: dict, entry_t: datetime.datetime,
                   exit_t: datetime.datetime,
                   entry_pos: tuple[float, float],
                   arc: list[tuple[datetime.datetime, float, float, float]],
                   pass_type: str = "overhead",
                   min_dist_deg: float = 0.0,
                   ) -> dict:
    duration_min = (exit_t - entry_t).total_seconds() / 60.0
    if arc:
        exit_lat, exit_lon = arc[-1][1], arc[-1][2]
        max_alt = max(p[3] for p in arc)
        dlat = arc[-1][1] - arc[0][1]
        dlon = arc[-1][2] - arc[0][2]
        if abs(dlat) > abs(dlon):
            direction = "N→S" if dlat < 0 else "S→N"
        else:
            direction = "W→E" if dlon > 0 else "E→W"
    else:
        exit_lat, exit_lon = entry_pos
        max_alt = 0.0
        direction = "N/A"

    nid = tle["norad_id"]
    # Sun elevation at the geometric midpoint of the crossing — gating
    # optical-EO arcs to daylight passes only is done by the caller.
    mid_t = entry_t + (exit_t - entry_t) / 2
    mid_lat = (entry_pos[0] + exit_lat) / 2
    mid_lon = (entry_pos[1] + exit_lon) / 2
    sun_elev = _solar_elevation_deg(mid_lat, mid_lon, mid_t)
    # Convert closest-approach degrees to approximate km for the report.
    # 1 deg latitude ~ 111 km. Slightly conservative on longitude.
    min_dist_km = round(min_dist_deg * 111.0, 0) if min_dist_deg else 0
    return {
        "norad_id":     nid,
        "name":         tle["name"],
        "country":      tle.get("country", ""),
        "operator":     tle.get("operator", ""),
        "sensor":       tle.get("sensor", "") or tle.get("type", ""),
        "resolution_m": tle.get("resolution_m"),
        "entry_utc":    entry_t.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "exit_utc":     exit_t.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "duration_min": round(duration_min, 1),
        "entry_lat":    round(entry_pos[0], 2),
        "entry_lon":    round(entry_pos[1], 2),
        "exit_lat":     round(exit_lat, 2),
        "exit_lon":     round(exit_lon, 2),
        "max_alt_km":   round(max_alt, 1),
        "direction":    direction,
        "sun_elev_deg": round(sun_elev, 1),
        "daytime":      sun_elev > DAYLIGHT_MIN_ELEV_DEG,
        "is_indian":    nid in INDIAN_NORAD_IDS,
        "pass_type":    pass_type,         # "overhead" | "tilt-range"
        "min_dist_km":  min_dist_km,        # 0 for overhead, >0 for tilt-range
        # Per-satellite imaging capability — used by Word reports and the
        # frontend map. Values come from hires_eo_satellites.TILT_SPECS.
        "altitude_km":      _alt_km(nid),
        "max_tilt_deg":     _tilt_deg(nid),
        "standoff_km":      _so_km(nid),
        # Strategic-site targets imaged by this pass (FEAT-004). A site is
        # "targeted" iff the satellite sub-point comes within its own
        # standoff_km of the site at some moment during the arc. Sorted
        # by closest approach. Empty list if nothing in range.
        "targeted_sites":   find_targeted_sites(arc, _so_km(nid)),
    }


# ── Observation / blind window computation ──────────────────────────────────
def _merge_observation_windows(crossings: list[dict],
                               day_start: datetime.datetime,
                               day_end:   datetime.datetime
                               ) -> tuple[list[dict], list[dict]]:
    """Merge per-satellite arcs into observation windows; invert to blind windows."""
    intervals = []
    for c in crossings:
        s = datetime.datetime.fromisoformat(c["entry_utc"].replace("Z", "+00:00"))
        e = datetime.datetime.fromisoformat(c["exit_utc"].replace("Z", "+00:00"))
        intervals.append((s, e, c["norad_id"]))

    intervals.sort(key=lambda iv: iv[0])

    merged: list[tuple[datetime.datetime, datetime.datetime, list[str]]] = []
    for s, e, nid in intervals:
        if merged and s <= merged[-1][1]:
            ps, pe, pn = merged[-1]
            merged[-1] = (ps, max(pe, e), pn + [nid])
        else:
            merged.append((s, e, [nid]))

    observation = [{
        "start":        s.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "end":          e.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "duration_min": round((e - s).total_seconds() / 60.0, 1),
        "sats":         list(dict.fromkeys(nids)),
    } for s, e, nids in merged]

    # Invert to find blind windows
    blind: list[dict] = []
    cursor = day_start
    for s, e, _ in merged:
        if s > cursor:
            blind.append(_blind_window(cursor, s))
        cursor = max(cursor, e)
    if cursor < day_end:
        blind.append(_blind_window(cursor, day_end))

    return observation, blind


def _blind_window(s: datetime.datetime, e: datetime.datetime) -> dict:
    return {
        "start":        s.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "end":          e.strftime("%Y-%m-%dT%H:%M:%SZ"),
        "duration_min": round((e - s).total_seconds() / 60.0, 1),
    }


# ── Orchestration ───────────────────────────────────────────────────────────
def build_forecast(start_date: datetime.date | None = None,
                   tles: dict[str, dict] | None = None
                   ) -> dict:
    if start_date is None:
        # Anchor the window to Pakistan's calendar date (PKT = UTC+5), so
        # "Day 1" is today in Pakistan — not today in UTC, which lags up to
        # 5h behind and made the chart start on "yesterday" for PK users.
        start_date = (datetime.datetime.now(datetime.timezone.utc)
                      + timedelta(hours=5)).date()
    # Space-Track top-up: pull the military / classified sats CelesTrak omits
    # (CSO, OFEQ, SAR-Lupe, SARah, …) straight into the TLE cache before we load
    # it. No-op unless the operator has added spacetrack_credentials.json.
    try:
        import spacetrack_tle
        from hires_eo_satellites import HIRES_EO_SATELLITES as _fleet
        _lut = {}
        for _s in _fleet:
            _lut[str(_s.norad_id)] = _s
            _lut[str(_s.norad_id).lstrip("0")] = _s
        spacetrack_tle.merge_into_forecast_cache(_lut)
    except Exception as _e:
        print(f"[SPACETRACK] top-up skipped: {_e}")
    if tles is None:
        tles = fetch_hires_eo_tles()

    t0 = time.time()
    print(f"[FORECAST] window: {start_date}  ->  {start_date + timedelta(days=FORECAST_DAYS-1)}")
    print(f"[FORECAST] propagator: {'sgp4 library' if _HAVE_SGP4 else 'hand-rolled fallback'}")

    indian_count = sum(1 for nid in tles if nid in INDIAN_NORAD_IDS)
    days_out: list[dict] = []

    for offset in range(FORECAST_DAYS):
        day = start_date + timedelta(days=offset)
        # PKT day boundaries (see _day_crossings): 00:00 PKT = 19:00 UTC prev day.
        day_start = (datetime.datetime(day.year, day.month, day.day, tzinfo=timezone.utc)
                     - timedelta(hours=5))
        day_end   = day_start + timedelta(days=1)

        day_crossings: list[dict] = []
        for tle in tles.values():
            day_crossings.extend(_day_crossings(tle, day))
        day_crossings.sort(key=lambda c: c["entry_utc"])

        obs, blind = _merge_observation_windows(day_crossings, day_start, day_end)

        # Helper — case-insensitive country match for Israel (the catalog
        # uses "Israel" but we don't want a typo bug to silently drop sats).
        def _is_israeli(c: dict) -> bool:
            return (c.get("country") or "").lower() == "israel"

        unique_sats     = len({c["norad_id"] for c in day_crossings})
        indian_cross    = sum(1 for c in day_crossings if c["is_indian"])
        ind_overhead    = sum(1 for c in day_crossings if c["is_indian"] and c["pass_type"] == "overhead")
        ind_tilt        = sum(1 for c in day_crossings if c["is_indian"] and c["pass_type"] == "tilt-range")
        israel_cross    = sum(1 for c in day_crossings if _is_israeli(c))
        isr_overhead    = sum(1 for c in day_crossings if _is_israeli(c) and c["pass_type"] == "overhead")
        isr_tilt        = sum(1 for c in day_crossings if _is_israeli(c) and c["pass_type"] == "tilt-range")
        # OTHER = not Indian AND not Israeli (Israeli now has its own bucket)
        oth_overhead    = sum(1 for c in day_crossings
                              if not c["is_indian"] and not _is_israeli(c)
                              and c["pass_type"] == "overhead")
        oth_tilt        = sum(1 for c in day_crossings
                              if not c["is_indian"] and not _is_israeli(c)
                              and c["pass_type"] == "tilt-range")
        overhead_total  = sum(1 for c in day_crossings if c["pass_type"] == "overhead")
        tilt_total      = sum(1 for c in day_crossings if c["pass_type"] == "tilt-range")
        blind_minutes   = round(sum(b["duration_min"] for b in blind), 1)
        longest_blind   = round(max((b["duration_min"] for b in blind), default=0), 1)

        # Split crossings by country bucket so consumers don't have to.
        india_crossings  = [c for c in day_crossings if c["is_indian"]]
        israel_crossings = [c for c in day_crossings if _is_israeli(c)]
        # OTHER now excludes Indian AND Israeli (Israel has its own report block)
        other_crossings  = [c for c in day_crossings
                            if not c["is_indian"] and not _is_israeli(c)]

        # Per-country blind windows — when is Pakistan invisible to a specific
        # country's satellite fleet. Same inversion logic, run per bucket.
        _, india_blind  = _merge_observation_windows(india_crossings,  day_start, day_end)
        _, israel_blind = _merge_observation_windows(israel_crossings, day_start, day_end)
        _, other_blind  = _merge_observation_windows(other_crossings,  day_start, day_end)
        india_blind_minutes  = round(sum(b["duration_min"] for b in india_blind),  1)
        israel_blind_minutes = round(sum(b["duration_min"] for b in israel_blind), 1)
        other_blind_minutes  = round(sum(b["duration_min"] for b in other_blind),  1)

        days_out.append({
            "date":                 str(day),
            "crossings":            day_crossings,         # all, for back-compat
            "india_crossings":      india_crossings,
            "israel_crossings":     israel_crossings,      # NEW — Israel split
            "other_crossings":      other_crossings,
            "observation_windows":  obs,
            "blind_windows":        blind,
            "india_blind_windows":  india_blind,
            "israel_blind_windows": israel_blind,           # NEW — Israeli-only
            "other_blind_windows":  other_blind,
            "totals": {
                "crossings":            len(day_crossings),
                "unique_sats":          unique_sats,
                "overhead_crossings":   overhead_total,
                "tilt_range_crossings": tilt_total,
                "indian_crossings":     indian_cross,
                "indian_overhead":      ind_overhead,
                "indian_tilt_range":    ind_tilt,
                "israeli_crossings":    israel_cross,       # NEW
                "israeli_overhead":     isr_overhead,       # NEW
                "israeli_tilt_range":   isr_tilt,           # NEW
                "other_overhead":       oth_overhead,
                "other_tilt_range":     oth_tilt,
                "blind_minutes":        blind_minutes,
                "longest_blind_min":    longest_blind,
                "india_blind_minutes":  india_blind_minutes,
                "israel_blind_minutes": israel_blind_minutes,  # NEW
                "other_blind_minutes":  other_blind_minutes,
            },
        })

        ind_marker = (f"  IN[OH:{ind_overhead} Tilt:{ind_tilt}]"
                      if indian_cross else "")
        print(f"[FORECAST] {day}: {overhead_total:>2d} overhead + "
              f"{tilt_total:>2d} tilt-range, {unique_sats:>2d} sats, "
              f"{blind_minutes:>5.0f} blind-min{ind_marker}")

    elapsed = time.time() - t0
    forecast = {
        "generated_at":           datetime.datetime.now(datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"),
        "start_date":             str(start_date),
        "end_date":                str(start_date + timedelta(days=FORECAST_DAYS - 1)),
        "pakistan_box": {"lat_min": PAK_LAT_MIN, "lat_max": PAK_LAT_MAX,
                          "lon_min": PAK_LON_MIN, "lon_max": PAK_LON_MAX},
        # Methodology version tag — bumped whenever the classifier or
        # targeting changes in a way that affects how passes are counted.
        # accuracy_tracker uses this to flag cross-version comparisons
        # (e.g. v1 archive vs v2 ground-truth) as methodology mismatches,
        # not SGP4-drift accuracy failures. History:
        #   v1 (pre-2026-05-13) — single 300 km global tilt buffer
        #   v2 (2026-05-13+)    — per-sat altitude * tan(max_tilt) standoff
        #                         + Pakistan strategic-site targeting
        "methodology_version":    "v2",
        "satellite_count":        len(tles),
        "indian_satellite_count": indian_count,
        "propagator":             "sgp4" if _HAVE_SGP4 else "handrolled",
        "duration_seconds":       round(elapsed, 1),
        "days":                   days_out,
    }

    _write_json_atomic(FORECAST_JSON, forecast)
    print(f"[FORECAST] wrote {FORECAST_JSON} in {elapsed:.1f}s")
    return forecast


def _write_json_atomic(path: str, payload: dict) -> None:
    tmp = path + ".tmp"
    with open(tmp, "w", encoding="utf-8") as f:
        json.dump(payload, f, indent=2)
        f.flush()
        os.fsync(f.fileno())
    os.replace(tmp, path)


# ── PKT time formatting (Pakistan Standard Time, UTC+5, no DST) ─────────────
# The operator reads these documents in Pakistan; UTC adds mental math.
# All time columns are rendered in PKT. The raw entry_utc/exit_utc fields
# in the JSON stay UTC for machine consumers — display layer converts.
def _pkt_time(iso_utc: str) -> str:
    """Convert 'YYYY-MM-DDTHH:MM:SSZ' to 'HH:MM PKT' (UTC + 5h, no DST)."""
    try:
        dt = datetime.datetime.strptime(iso_utc, "%Y-%m-%dT%H:%M:%SZ")
    except (ValueError, TypeError):
        return iso_utc[11:19] if iso_utc else "-"
    return (dt + timedelta(hours=5)).strftime("%H:%M")


def _pkt_datetime(iso_utc: str) -> str:
    """Full 'YYYY-MM-DD HH:MM PKT' for ranges that may cross midnight."""
    try:
        dt = datetime.datetime.strptime(iso_utc, "%Y-%m-%dT%H:%M:%SZ")
    except (ValueError, TypeError):
        return iso_utc
    return (dt + timedelta(hours=5)).strftime("%Y-%m-%d %H:%M")


# ── Optical vs SAR split helper ─────────────────────────────────────────────
# The operator reads two distinct threat tracks: visible-light cameras (need
# daylight, see colour/shape) and synthetic-aperture radar (any time, sees
# through cloud and dark). The report keeps them in separate, labelled
# sub-sections so the two can be compared without scanning a mixed list.
def _split_optical_sar(crossings: list[dict]) -> tuple[list[dict], list[dict]]:
    """Return (optical_list, sar_list). A pass is SAR iff its sensor string
    contains 'SAR' (case-insensitive); everything else (including
    Optical-Military like CSO) is optical."""
    optical, sar = [], []
    for c in crossings:
        if "SAR" in (c.get("sensor", "") or "").upper():
            sar.append(c)
        else:
            optical.append(c)
    return optical, sar


# ── Word renderer ───────────────────────────────────────────────────────────
def _add_crossings_table(doc, crossings, accent_color, Pt, RGBColor, indian: bool):
    """Render a single per-country crossings table; tilt-range rows shaded.

    Columns now include per-satellite imaging capability — Max Tilt and
    Standoff km — so the reader can immediately see what each sat's reach
    is rather than assuming a one-size-fits-all 300 km buffer.
    """
    AMBER = RGBColor(0xE6, 0x7E, 0x22)   # tilt-range marker
    TARGET_RED = RGBColor(0xC8, 0x10, 0x2E)
    cols = ["Time PKT", "Pass", "Satellite", "Country",
            "Res (m)", "Tilt (°)", "Standoff (km)",
            "Off-nadir (km)", "Duration", "Direction", "Targets (PK)"]
    tbl = doc.add_table(rows=1 + len(crossings), cols=len(cols))
    tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(cols):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    for i, c in enumerate(crossings, 1):
        row = tbl.rows[i].cells
        pass_label = "OVER" if c["pass_type"] == "overhead" else "TILT"
        row[0].text = _pkt_time(c["entry_utc"])
        row[1].text = pass_label
        row[2].text = c["name"][:30]
        row[3].text = c["country"]
        row[4].text = f"{c['resolution_m']:.2f}" if c.get("resolution_m") else "-"
        row[5].text = f"{c['max_tilt_deg']}" if c.get("max_tilt_deg") is not None else "-"
        row[6].text = f"{c['standoff_km']:.0f}" if c.get("standoff_km") is not None else "-"
        # Off-nadir distance from box edge — 0 if overhead, else km
        row[7].text = "0" if c["pass_type"] == "overhead" else f"~{int(c.get('min_dist_km', 0))}"
        row[8].text = f"{c['duration_min']:.1f} min"
        row[9].text = c["direction"]
        # Targets — UNIQUE city list, tier-1 first. Empty cell when the
        # pass doesn't bring any catalog site inside the sat's reach.
        # Dedupe by city so "Rawalpindi, Rawalpindi" (GHQ + Nur Khan both
        # in Rawalpindi) collapses to one entry — readable for the operator.
        tgt_list = c.get("targeted_sites") or []
        tier1 = [t for t in tgt_list if t.get("tier") == 1]
        tier2 = [t for t in tgt_list if t.get("tier") == 2]
        seen = set(); t1_cities = []
        for t in tier1:
            ct = t.get("city")
            if ct and ct not in seen:
                seen.add(ct); t1_cities.append(ct)
        seen2 = set(); t2_cities = []
        for t in tier2:
            ct = t.get("city")
            if ct and ct not in seen2:
                seen2.add(ct); t2_cities.append(ct)
        if t1_cities or t2_cities:
            primary = ", ".join(t1_cities) if t1_cities else ", ".join(t2_cities[:4])
            extra = f"  +{len(t2_cities)}" if (t1_cities and t2_cities) else ""
            row[10].text = primary + extra
        else:
            row[10].text = "—"
        for cell in row:
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8)
        # Highlight: Indian sats use the accent (red); tilt-range row uses amber
        if indian:
            for cell in (row[2], row[3]):
                for r in cell.paragraphs[0].runs:
                    r.font.color.rgb = accent_color
                    r.bold = True
        if c["pass_type"] == "tilt-range":
            for r in row[1].paragraphs[0].runs:
                r.font.color.rgb = AMBER
                r.bold = True
        # Bold-red the Targets cell if any tier-1 site is reachable on this pass
        if tier1:
            for r in row[10].paragraphs[0].runs:
                r.font.color.rgb = TARGET_RED
                r.bold = True


_COVERAGE_CATEGORIES = {
    "combined": {
        "title":      "1) COMBINED (SAR + OPTICAL) – ANY INDIAN SATELLITE COVERAGE",
        "bar_rgb":    (91, 45, 140),
        "border_rgb": (138, 95, 195),
        "predicate":  lambda c: bool(c.get("is_indian")),
    },
    "optical": {
        "title":      "2) OPTICAL ONLY (DAYLIGHT SATELLITES)",
        "bar_rgb":    (184, 134, 11),
        "border_rgb": (224, 174, 51),
        "predicate":  lambda c: bool(c.get("is_indian"))
                                and "SAR" not in str(c.get("sensor", "")).upper(),
    },
    "sar": {
        "title":      "3) SAR ONLY (ALL-WEATHER, DAY + NIGHT)",
        "bar_rgb":    (31, 58, 95),
        "border_rgb": (88, 138, 195),
        "predicate":  lambda c: bool(c.get("is_indian"))
                                and "SAR" in str(c.get("sensor", "")).upper(),
    },
}


def generate_coverage_pdf(forecast_json_path: str | None = None,
                          out_pdf_path: str | None = None,
                          country: str = "india",
                          n_days: int = 15) -> str | None:
    """Re-render the category PNGs (Combined / Optical / SAR) for the
    requested country bucket from the latest forecast JSON on disk and
    bundle them into a plotter-ready multipage landscape PDF. Returns the
    PDF path on success, None on failure.

    Parameters
    ----------
    country : ``"india"`` (default) / ``"china"`` / ``"other"``
    n_days  : 5 / 10 / 15 — number of forecast days to render (truncates
              the forecast's 15-day window if smaller).

    Wired to the ATLAS dashboard's three crossing-chart buttons so the
    operator can produce a print-ready file on demand without rebuilding
    the whole Word report.
    """
    import json
    fj = forecast_json_path or FORECAST_JSON
    if not os.path.isfile(fj):
        print(f"[PDF] forecast JSON not found at {fj}")
        return None
    with open(fj, "r", encoding="utf-8") as f:
        forecast = json.load(f)

    country = (country or "india").lower()
    tag = country.capitalize()        # India / China / Other — used in filenames
    n_days = max(1, min(15, int(n_days or 15)))
    cat_files = [
        ("combined", f"Pakistan_{n_days}Day_Coverage_{tag}_Combined.png"),
        ("optical",  f"Pakistan_{n_days}Day_Coverage_{tag}_Optical.png"),
        ("sar",      f"Pakistan_{n_days}Day_Coverage_{tag}_SAR.png"),
    ]
    # OTHER-country chart drops the Combined page — the merged window list
    # is too dense to be useful with ~25 contributing satellites, and the
    # Optical/SAR pages already convey the same information split cleanly.
    if country == "other":
        cat_files = [t for t in cat_files if t[0] != "combined"]
    png_paths: list[str] = []
    for cat, fname in cat_files:
        png_path = os.path.join(DOCS_DIR, fname)
        try:
            _render_category_png(forecast, cat, png_path,
                                 country=country, n_days=n_days)
            png_paths.append(png_path)
            print(f"[PDF] rendered {fname}")
        except Exception as e:
            print(f"[PDF] render {cat} failed: {e}")

    if not png_paths:
        return None

    if out_pdf_path is None:
        out_pdf_path = os.path.join(
            DOCS_DIR, f"Pakistan_{n_days}Day_Coverage_{tag}.pdf")
    try:
        from PIL import Image
        imgs = []
        for _p in png_paths:
            _src = Image.open(_p); imgs.append(_src.convert("RGB")); _src.close()
        # Embed at 1600 px / 14.4 in ~= 111.11 dpi so each PDF page is
        # exactly 1.2 ft (14.4") wide x 9" tall on the plotter.
        imgs[0].save(out_pdf_path, format="PDF", resolution=1600 / 14.4,
                     save_all=True, append_images=imgs[1:])
        print(f"[PDF] wrote {out_pdf_path}")
        _cleanup_pngs(png_paths)       # PDF has them embedded — drop the PNGs
        return out_pdf_path
    except Exception as e:
        print(f"[PDF] save failed: {e}")
        return None


def generate_observation_timeline_pdf(forecast_json_path: str | None = None,
                                      out_pdf_path: str | None = None,
                                      n_days: int = 15) -> str | None:
    """ALL-COUNTRY observation timeline — one horizontal 24-hour bar per day,
    WHITE = observing (a satellite overhead) / GREEN = non-observing (blind,
    safe). Produces a 3-PAGE PDF:
        page 1 — OPTICAL + SAR  (any satellite)
        page 2 — OPTICAL only   (daylight imagers)
        page 3 — SAR only       (all-weather radar)
    Each page merges every country; coverage is recomputed per page from the
    matching crossings. Plotter-ready 14.4" x 9" pages."""
    import json
    from PIL import Image, ImageDraw, ImageFont
    import datetime as _dt

    fj = forecast_json_path or FORECAST_JSON
    if not os.path.isfile(fj):
        print(f"[PDF] forecast JSON not found at {fj}")
        return None
    with open(fj, "r", encoding="utf-8") as f:
        forecast = json.load(f)
    n_days = max(1, min(15, int(n_days or 15)))
    days = forecast.get("days", [])[:n_days]
    if not days:
        print("[PDF] no forecast days to render")
        return None

    W, H = 1600, 1000
    BG     = (247, 249, 252)
    BAR_BG = (228, 233, 240)
    TITLE  = (22, 28, 42)
    SUB    = (92, 100, 114)
    CYAN   = (0, 122, 156)
    GOLD   = (158, 106, 0)
    GREEN  = (37, 152, 79)     # non-observing (safe)
    WHITE  = (255, 255, 255)   # observing (satellite overhead)
    GRID   = (150, 164, 182)
    BORDER = (52, 64, 82)
    DAYLBL = (38, 102, 132)

    def _font(name, size):
        try:
            return ImageFont.truetype(rf"C:\Windows\Fonts\{name}", size)
        except OSError:
            return ImageFont.load_default()
    F_TITLE = _font("arialbd.ttf", 24)
    F_SUB   = _font("arial.ttf", 15)
    F_SUBB  = _font("arialbd.ttf", 15)
    F_BANNER= _font("arialbd.ttf", 18)
    F_DAY   = _font("arialbd.ttf", 13)
    F_DATE  = _font("arial.ttf", 11)
    F_TICK  = _font("consolab.ttf", 12)
    F_LEG   = _font("arialbd.ttf", 13)
    F_FOOT  = _font("arial.ttf", 10)
    F_FOOTB = _font("arialbd.ttf", 11)

    sd = forecast.get("start_date", "?")
    ed = days[-1].get("date", forecast.get("end_date", "?"))
    try:
        sd_h = _dt.datetime.strptime(sd, "%Y-%m-%d").strftime("%d %b %Y")
        ed_h = _dt.datetime.strptime(ed, "%Y-%m-%d").strftime("%d %b %Y")
    except ValueError:
        sd_h, ed_h = sd, ed

    def _ufrac(iso):
        t = _dt.datetime.strptime(iso, "%Y-%m-%dT%H:%M:%SZ")
        return (t.hour * 3600 + t.minute * 60 + t.second) / 86400.0

    def _sensor_pred(cat):
        if cat == "optical":
            return lambda c: "SAR" not in str(c.get("sensor", "")).upper()
        if cat == "sar":
            return lambda c: "SAR" in str(c.get("sensor", "")).upper()
        return lambda c: True

    def _obs_windows(day, pred):
        """Merge the matching crossings into (start_frac, end_frac, minutes)
        observing windows for this day."""
        rel = sorted((c for c in day.get("crossings", []) if pred(c)),
                     key=lambda c: c["entry_utc"])
        if not rel:
            return []
        try:
            merged = [[_dt.datetime.strptime(rel[0]["entry_utc"], "%Y-%m-%dT%H:%M:%SZ"),
                       _dt.datetime.strptime(rel[0]["exit_utc"],  "%Y-%m-%dT%H:%M:%SZ")]]
            for c in rel[1:]:
                cs = _dt.datetime.strptime(c["entry_utc"], "%Y-%m-%dT%H:%M:%SZ")
                ce = _dt.datetime.strptime(c["exit_utc"],  "%Y-%m-%dT%H:%M:%SZ")
                if cs <= merged[-1][1]:
                    if ce > merged[-1][1]:
                        merged[-1][1] = ce
                else:
                    merged.append([cs, ce])
        except (ValueError, KeyError):
            return []
        out = []
        for s, e in merged:
            sf = (s.hour*3600 + s.minute*60 + s.second) / 86400.0
            ef = (e.hour*3600 + e.minute*60 + e.second) / 86400.0
            if ef > sf:
                out.append((sf, ef, (e - s).total_seconds() / 60.0))
        return out

    # ── One page per sensor category ────────────────────────────────────────
    CATS = [
        ("combined", "1)  OPTICAL + SAR  —  ANY SATELLITE OVERHEAD", (96, 46, 134)),
        ("optical",  "2)  OPTICAL ONLY  —  daylight imagers",        CYAN),
        ("sar",      "3)  SAR ONLY  —  all-weather radar (day & night)", GOLD),
    ]

    def _render_page(cat, cat_title, cat_color):
        img = Image.new("RGB", (W, H), BG)
        d = ImageDraw.Draw(img)
        pred = _sensor_pred(cat)

        # Header
        d.rectangle([0, 0, W, 64], fill=BAR_BG)
        d.text((24, 10),
               f"{n_days}-DAY ALL-COUNTRY OBSERVATION TIMELINE OVER PAKISTAN (PKT)",
               font=F_TITLE, fill=TITLE)
        x = 26
        d.text((x, 40), "Forecast Window: ", font=F_SUB, fill=SUB)
        x += d.textlength("Forecast Window: ", font=F_SUB)
        d.text((x, 40), f"{sd_h} – {ed_h}", font=F_SUBB, fill=CYAN)
        x += d.textlength(f"{sd_h} – {ed_h}", font=F_SUBB)
        d.text((x, 40), "    |    Pakistan Standard Time (PKT, UTC+5)",
               font=F_SUB, fill=SUB)

        # Category banner
        by = 70
        d.rectangle([24, by, W - 24, by + 30], fill=cat_color)
        d.text((34, by + 6), cat_title, font=F_BANNER, fill=(255, 255, 255))

        # Legend
        ly = by + 40
        d.rectangle([24, ly, 48, ly + 18], fill=WHITE, outline=BORDER, width=1)
        d.text((54, ly + 1), "Observing — satellite overhead", font=F_LEG, fill=TITLE)
        x2 = 54 + d.textlength("Observing — satellite overhead", font=F_LEG) + 36
        d.rectangle([x2, ly, x2 + 24, ly + 18], fill=GREEN, outline=BORDER, width=1)
        d.text((x2 + 30, ly + 1), "Non-observing — blind window (safe)",
               font=F_LEG, fill=TITLE)

        # Grid geometry
        gx = 24
        gy = ly + 30
        lbl_w = 150
        gw = W - 2 * gx
        tl_x0 = gx + lbl_w
        tl_w = gw - lbl_w
        grid_bottom = H - 72
        axis_h = 22
        rows_y0 = gy + axis_h
        row_h = (grid_bottom - rows_y0) / len(days)

        # Hourly labels + short ticks in the axis strip (the full-height
        # 15-minute grid is overlaid AFTER the bars, below).
        d.text((gx + 8, gy + 4), "PKT", font=F_TICK, fill=SUB)
        for hh in range(0, 25):
            xx = tl_x0 + tl_w * (hh / 24.0)
            d.line([(xx, gy + axis_h - 6), (xx, gy + axis_h)], fill=GRID, width=1)
            lab = f"{(hh + 5) % 24:02d}"
            d.text((xx - d.textlength(lab, font=F_TICK) / 2, gy + 4),
                   lab, font=F_TICK, fill=GOLD)

        cov_minutes = []
        for i, day in enumerate(days):
            ry0 = rows_y0 + i * row_h
            ry1 = ry0 + row_h
            d.rectangle([gx, ry0, tl_x0, ry1], fill=(231, 236, 242))
            d.text((gx + 8, ry0 + 5), f"DAY {i+1}", font=F_DAY, fill=DAYLBL)
            try:
                do = _dt.datetime.strptime(day["date"], "%Y-%m-%d")
                d.text((gx + 8, ry0 + 22), do.strftime("%d %b  %a"),
                       font=F_DATE, fill=GOLD)
            except (ValueError, KeyError):
                pass
            bt, bb = ry0 + 5, ry1 - 5
            wins = _obs_windows(day, pred)
            covmin = sum(m for _, _, m in wins)
            d.rectangle([tl_x0, bt, tl_x0 + tl_w, bb], fill=GREEN)
            # Exact (continuous) observing windows — the 15-minute grid is
            # overlaid below for reading times, without coarsening the data.
            for sf, ef, _ in wins:
                d.rectangle([tl_x0 + tl_w * sf, bt, tl_x0 + tl_w * ef, bb],
                            fill=WHITE)
            cov_minutes.append(covmin)
            d.rectangle([tl_x0, bt, tl_x0 + tl_w, bb], outline=BORDER, width=1)
            d.line([(gx, ry1), (gx + gw, ry1)], fill=GRID, width=1)

        # 15-minute grid overlaid on the bars — faint every 15 min, darker on
        # the hour — so the timeline reads at 15-minute resolution.
        MINOR = (214, 221, 229)
        for q in range(0, 97):
            xx = tl_x0 + tl_w * (q / 96.0)
            d.line([(xx, rows_y0), (xx, grid_bottom)],
                   fill=(GRID if q % 4 == 0 else MINOR), width=1)

        d.rectangle([gx, gy, gx + gw, grid_bottom], outline=BORDER, width=2)

        # Footer
        fy = H - 62
        d.rectangle([0, fy, W, H], fill=BAR_BG)
        d.text((12, fy + 6), "NOTES:", font=F_FOOTB, fill=TITLE)
        d.text((12, fy + 22),
               "White = observing (a matching satellite overhead).  "
               "Green = non-observing (blind window — safe).  "
               "Grid = 15-minute intervals.",
               font=F_FOOT, fill=SUB)
        cat_note = {"combined": "Optical + SAR, merged across ALL countries "
                                "(India · China · USA · Israel · others).",
                    "optical":  "Optical / daylight imagers only, all countries. "
                                "Capture only in daylight.",
                    "sar":      "SAR / radar only, all countries. Capture day & "
                                "night, through cloud."}[cat]
        d.text((12, fy + 38), cat_note + "  Generated from your 15-day forecast.",
               font=F_FOOT, fill=SUB)
        avg_cov = sum(cov_minutes) / len(cov_minutes) if cov_minutes else 0
        avg_blind = 1440 - avg_cov
        bx = 1080
        d.rectangle([bx, fy + 6, bx + 360, fy + 54], outline=cat_color, width=1)
        d.text((bx + 12, fy + 9), "DAILY AVERAGE", font=F_FOOTB, fill=cat_color)
        d.text((bx + 12, fy + 25), f"Observing:      ~ {avg_cov:.0f} min/day",
               font=F_FOOT, fill=SUB)
        d.text((bx + 12, fy + 39),
               f"Non-observing:  ~ {int(avg_blind // 60)} h {int(avg_blind % 60)} m/day",
               font=F_FOOT, fill=SUB)
        when = (_dt.datetime.now(_dt.timezone.utc) + _dt.timedelta(hours=5)
                ).strftime("%d %b %Y %H:%M PKT")
        d.text((W - 250, fy + 18), "Generated on:", font=F_FOOT, fill=SUB)
        d.text((W - 250, fy + 34), when, font=F_FOOTB, fill=TITLE)
        return img

    try:
        os.makedirs(DOCS_DIR, exist_ok=True)
        pages = []
        for cat, title, color in CATS:
            page = _render_page(cat, title, color)
            page.save(os.path.join(
                DOCS_DIR, f"Pakistan_{n_days}Day_Observation_Timeline_{cat.capitalize()}.png"))
            pages.append(page.convert("RGB"))
        if out_pdf_path is None:
            out_pdf_path = os.path.join(
                DOCS_DIR, f"Pakistan_{n_days}Day_Observation_Timeline.pdf")
        pages[0].save(out_pdf_path, format="PDF", resolution=1600 / 14.4,
                      save_all=True, append_images=pages[1:])
        print(f"[PDF] wrote {out_pdf_path}  (3 pages: combined / optical / sar)")
        return out_pdf_path
    except Exception as e:
        print(f"[PDF] timeline save failed: {e}")
        return None


_COUNTRY_LABELS = {
    "india":  "INDIAN",
    "china":  "CHINESE",
    "usa":    "USA",
    "israel": "ISRAELI",
    "other":  "OTHER-COUNTRY",
    "all":    "ALL-COUNTRY",
}


def _country_match(crossing: dict, country: str) -> bool:
    """True when a crossing's country bucket matches the requested chart.
    ``other`` is everything that isn't India, China, USA, or Israel — those
    four each have their own dedicated chart."""
    country = country.lower()
    c_field = (crossing.get("country") or "").lower()
    if country == "all":
        return True                       # every country combined
    if country == "india":
        return bool(crossing.get("is_indian"))
    if country == "china":
        return c_field == "china"
    if country == "usa":
        return c_field == "usa"
    if country == "israel":
        return c_field == "israel"
    if country == "other":
        return (not bool(crossing.get("is_indian"))) \
               and c_field != "china" \
               and c_field != "usa" \
               and c_field != "israel"
    return False


def _render_category_png(forecast: dict, category: str, out_png: str,
                        country: str = "india",
                        n_days: int = 15) -> str:
    """Render a single-category coverage timetable as a pixel-exact wall-
    planner PNG, filtered to the requested country bucket
    (india / china / other) and trimmed to the first ``n_days`` (5/10/15).
    The full landscape page hosts ONE coverage category
    (Combined / Optical / SAR), with a soft neon glow around the green
    time-window text. Returns the PNG path."""
    from PIL import Image, ImageDraw, ImageFont, ImageFilter
    import datetime as _dt
    _PKT = _dt.timedelta(hours=5)
    cfg = _COVERAGE_CATEGORIES[category]
    country = country.lower()
    country_label = _COUNTRY_LABELS.get(country, country.upper())
    n_days = max(1, min(15, int(n_days)))
    # ALL-COUNTRY merges ~500 satellites — collapse near-adjacent coverage
    # (<= 15 min apart) into single blocks so cells stay readable. Per-country
    # charts keep exact per-pass windows (gap 0).
    _GAP = 15 if country == "all" else 0

    # ── Canvas + palette (light print theme — plotter-ready 14.4" x 9") ─────
    # Canvas stays 1600x1000 px (aspect 1.6:1). generate_coverage_pdf()
    # embeds it at 1600/14.4 ~= 111.11 dpi so the PDF page is exactly
    # 1.2 ft (14.4") wide x 9" tall for the plotter.
    W, H = 1600, 1000
    BG          = (247, 249, 252)   # near-white page background
    BAR_BG      = (228, 233, 240)   # light header / footer strip
    TITLE       = (22, 28, 42)      # near-black primary text
    SUB_GREY    = (92, 100, 114)    # medium-grey secondary text
    CYAN        = (0, 122, 156)     # deep teal accent
    GOLD        = (158, 106, 0)     # dark amber accent
    AMBER       = (158, 106, 0)     # dark amber (day-header dates)
    CELL_BG     = (255, 255, 255)   # white grid cell
    GREEN_TXT   = (18, 120, 52)     # deep green — coverage windows
    DAY_LABEL   = (38, 102, 132)    # deep blue — DAY N labels
    BLIND_RED   = (200, 55, 55)     # red — blind-window marker

    img  = Image.new("RGBA", (W, H), BG + (255,))
    draw = ImageDraw.Draw(img)
    # Glow layer kept for structural symmetry but unused in the light
    # theme — a neon halo doesn't read on a white page, so coverage text
    # is drawn sharp. This empty layer composites as a harmless no-op.
    glow  = Image.new("RGBA", (W, H), (0, 0, 0, 0))
    gdraw = ImageDraw.Draw(glow)
    # Re-draw queue so each green-time string is painted sharp on top after
    # the glow pass has been blurred + composited under it.
    sharp_queue: list[tuple[tuple[int, int], str]] = []

    def _font(name: str, size: int):
        for base in (r"C:\Windows\Fonts", "/usr/share/fonts/truetype/dejavu",
                     "/Library/Fonts", "/System/Library/Fonts"):
            try:
                return ImageFont.truetype(f"{base}/{name}", size)
            except (OSError, FileNotFoundError):
                continue
        return ImageFont.load_default()

    F_TITLE    = _font("arialbd.ttf", 26)
    F_SUB      = _font("arial.ttf", 15)
    F_SUB_B    = _font("arialbd.ttf", 15)
    F_LEG_HDR  = _font("arialbd.ttf", 12)
    F_LEG      = _font("arial.ttf", 10)
    F_SECTION  = _font("arialbd.ttf", 18)
    F_DAY      = _font("arialbd.ttf", 13)
    F_DATE     = _font("arial.ttf", 11)
    F_TIME     = _font("consolab.ttf", 18)   # left time-label column (HH:00)
    F_CELL     = _font("consolab.ttf", 14)   # in-grid window text — bold + bigger
    F_FOOT_HDR = _font("arialbd.ttf", 11)
    F_FOOT     = _font("arial.ttf", 10)

    # ── Top header strip ────────────────────────────────────────────────────
    draw.rectangle([0, 0, W, 70], fill=BAR_BG)

    # Tiny satellite glyph (drawn, not loaded)
    sx, sy = 20, 18
    draw.rectangle([sx+14, sy+12, sx+28, sy+26], fill=CYAN)
    draw.rectangle([sx,    sy+15, sx+12, sy+23], outline=CYAN, width=1)
    draw.rectangle([sx+30, sy+15, sx+42, sy+23], outline=CYAN, width=1)
    for r in (8, 14):
        draw.arc([sx+21-r, sy-r, sx+21+r, sy+r], 220, 320, fill=CYAN, width=2)

    draw.text((78, 10),
              f"{n_days}-DAY {country_label} SATELLITE COVERAGE TIMETABLE OVER PAKISTAN (PKT)",
              font=F_TITLE, fill=TITLE)

    sd = forecast.get("start_date", "?")
    # End date follows the (possibly truncated) day list, not the JSON's
    # original 15-day window.
    _days_for_window = forecast.get("days", [])[:n_days]
    ed = _days_for_window[-1].get("date") if _days_for_window else forecast.get("end_date", "?")
    try:
        sd_h = _dt.datetime.strptime(sd, "%Y-%m-%d").strftime("%d %b %Y")
        ed_h = _dt.datetime.strptime(ed, "%Y-%m-%d").strftime("%d %b %Y")
    except ValueError:
        sd_h, ed_h = sd, ed
    x, y = 78, 44
    draw.text((x, y), "Forecast Window: ", font=F_SUB, fill=SUB_GREY)
    x += draw.textlength("Forecast Window: ", font=F_SUB)
    draw.text((x, y), f"{sd_h} – {ed_h}", font=F_SUB_B, fill=CYAN)
    x += draw.textlength(f"{sd_h} – {ed_h}", font=F_SUB_B)
    draw.text((x, y), "    |    Time Zone: ", font=F_SUB, fill=SUB_GREY)
    x += draw.textlength("    |    Time Zone: ", font=F_SUB)
    draw.text((x, y), "Pakistan Standard Time (PKT, UTC+5)",
              font=F_SUB_B, fill=GOLD)

    # (top-right corner intentionally left blank)

    # ── Helpers: merge crossings → PKT HH:MM – HH:MM windows ───────────────
    # gap_min > 0 merges windows separated by <= gap_min minutes. Used by the
    # ALL-COUNTRY chart so ~500 satellites' overlapping/near coverage collapses
    # into a few readable blocks instead of dozens of colliding entries.
    def _windows(day, pred, gap_min: int = 0):
        rel = sorted((c for c in day.get("crossings", []) if pred(c)),
                     key=lambda c: c["entry_utc"])
        if not rel:
            return []
        gap = _dt.timedelta(minutes=gap_min)
        def _p(s):
            return _dt.datetime.strptime(s, "%Y-%m-%dT%H:%M:%SZ")
        try:
            m = [[_p(rel[0]["entry_utc"]), _p(rel[0]["exit_utc"])]]
            for c in rel[1:]:
                cs, ce = _p(c["entry_utc"]), _p(c["exit_utc"])
                if cs <= m[-1][1] + gap:           # overlapping or within gap
                    if ce > m[-1][1]:
                        m[-1][1] = ce
                else:
                    m.append([cs, ce])
        except (ValueError, KeyError):
            return []
        out = []
        for s, e in m:
            try:
                ps = s + _PKT
                pe = e + _PKT
                # Compact format: same hour → "HH:MM-MM", crossing hour →
                # "HH:MM-HH:MM" (no surrounding spaces). Keeps every entry
                # narrow enough to fit the day column at the larger 14 pt
                # cell font without spilling into neighbouring cells.
                if ps.hour == pe.hour:
                    out.append(f"{ps.strftime('%H:%M')}-{pe.strftime('%M')}")
                else:
                    out.append(f"{ps.strftime('%H:%M')}-{pe.strftime('%H:%M')}")
            except (ValueError, KeyError):
                pass
        return out

    # Build a country+category predicate dynamically. The static predicates
    # in _COVERAGE_CATEGORIES were Indian-only; here we re-apply the
    # category's sensor filter on top of the requested country bucket.
    def pred(c: dict) -> bool:
        if not _country_match(c, country):
            return False
        sensor = str(c.get("sensor", "")).upper()
        if category == "combined":
            return True
        if category == "optical":
            return "SAR" not in sensor
        if category == "sar":
            return "SAR" in sensor
        return False

    # Section title — rewrite the COMBINED line to reflect the country
    # bucket; OPTICAL/SAR titles are country-independent.
    if category == "combined":
        section_title = (f"1) COMBINED (SAR + OPTICAL) – "
                         f"ANY {country_label} SATELLITE COVERAGE")
    else:
        section_title = cfg["title"]

    days = forecast.get("days", [])[:n_days]
    n_cols = max(1, len(days))   # one column per day actually present

    # ── Class-timetable layout: rows=hours of day, columns=n_days ──────────
    sec_y    = 80
    side_pad = 12
    header_h = 38

    # Section header bar
    draw.rectangle([side_pad, sec_y, W - side_pad, sec_y + header_h],
                   fill=cfg["bar_rgb"])
    draw.text((22, sec_y + 9), section_title, font=F_SECTION, fill=(255, 255, 255))

    # Per-day windows, bucketed by start-hour: List[Dict[hour, List[str]]]
    day_hour_windows = []
    active_hours: set[int] = set()
    for d in days:
        wins = _windows(d, pred, gap_min=_GAP)
        hmap: dict[int, list[str]] = {}
        for w in wins:
            try:
                hh = int(w.split(":")[0])
                hmap.setdefault(hh, []).append(w)
                active_hours.add(hh)
            except (ValueError, IndexError):
                pass
        day_hour_windows.append(hmap)

    # If a category has no passes at all, still show a couple of placeholder
    # rows so the timetable isn't empty.
    hours = sorted(active_hours) if active_hours else [9, 10, 11]

    grid_x      = side_pad
    grid_y      = sec_y + header_h + 6
    grid_w      = W - 2 * side_pad
    time_col_w  = 90
    day_col_w   = (grid_w - time_col_w) // n_cols
    grid_max_y  = H - 70                       # leave room for footer strip
    body_avail  = grid_max_y - (grid_y + 62)   # 62 = day-header row height (DAY N · date · weekday)
    GRID_LINE   = (184, 192, 203)   # light grid lines
    ROW_ALT     = (240, 243, 247)   # alternating row tint
    ROW_BASE    = (255, 255, 255)   # base row (white)
    TIME_BG     = (231, 236, 242)   # time column / day-header strip
    TIME_FG     = (124, 82, 0)      # dark amber time labels

    # Per-hour row heights — sized to the busiest cell in that hour only,
    # so sparse rows (e.g. 05:00 / 06:00) collapse and dense rows (10:00 /
    # 11:00) stretch. 14 pt cell font = 19 px per stacked window line +
    # 20 px top/bottom padding. Minimum 38 px even for empty rows.
    row_heights: dict[int, int] = {}
    for _h in hours:
        _max_in_row = max(
            (len(hmap.get(_h, [])) for hmap in day_hour_windows),
            default=0,
        )
        row_heights[_h] = max(38, 20 + 19 * max(1, _max_in_row))
    # If the column of rows would overflow the available vertical space,
    # scale them all down proportionally so they still fit on one page.
    _total_rows = sum(row_heights.values())
    if _total_rows > body_avail and _total_rows > 0:
        _scale = body_avail / _total_rows
        row_heights = {h: max(30, int(rh * _scale))
                       for h, rh in row_heights.items()}

    # Day-header row (DAY N · date · weekday)
    # Three stacked text lines per day column:
    #   line 1 (+4)  : "DAY 1"            — F_DAY,  DAY_LABEL  (deep blue, bold)
    #   line 2 (+22) : "25 May"           — F_DATE, AMBER      (orange)
    #   line 3 (+40) : "Sunday"           — F_DATE, DAY_LABEL  (deep blue, regular)
    hdr_y, hdr_h = grid_y, 62
    draw.rectangle([grid_x, hdr_y, grid_x + grid_w, hdr_y + hdr_h], fill=TIME_BG)
    # TIME label
    tlbl = "TIME (PKT)"
    tw = draw.textlength(tlbl, font=F_DAY)
    draw.text((grid_x + (time_col_w - tw)//2, hdr_y + (hdr_h - 16)//2),
              tlbl, font=F_DAY, fill=TIME_FG)
    draw.line([(grid_x + time_col_w, hdr_y),
               (grid_x + time_col_w, hdr_y + hdr_h)], fill=GRID_LINE, width=1)
    for di in range(n_cols):
        cx = grid_x + time_col_w + di * day_col_w
        day_lbl = f"DAY {di+1}"
        tw = draw.textlength(day_lbl, font=F_DAY)
        draw.text((cx + (day_col_w - tw)//2, hdr_y + 4),
                  day_lbl, font=F_DAY, fill=DAY_LABEL)
        if di < len(days):
            try:
                _dt_obj = _dt.datetime.strptime(days[di]["date"], "%Y-%m-%d")
                ds  = _dt_obj.strftime("%d %b")
                # Full weekday if it fits in the column at this font size,
                # otherwise fall back to the 3-letter abbreviation. Keeps
                # narrow 15-day plots readable without clipping.
                full = _dt_obj.strftime("%A")          # e.g. "Sunday"
                abbr = _dt_obj.strftime("%a")          # e.g. "Sun"
                wname = full if draw.textlength(full, font=F_DATE) <= day_col_w - 6 else abbr
            except (ValueError, KeyError):
                ds, wname = days[di].get("date", "")[:6], ""
            # Date (orange)
            tw = draw.textlength(ds, font=F_DATE)
            draw.text((cx + (day_col_w - tw)//2, hdr_y + 22),
                      ds, font=F_DATE, fill=AMBER)
            # Weekday name (deep blue, third line)
            if wname:
                tw = draw.textlength(wname, font=F_DATE)
                draw.text((cx + (day_col_w - tw)//2, hdr_y + 42),
                          wname, font=F_DATE, fill=DAY_LABEL)
        draw.line([(cx + day_col_w, hdr_y),
                   (cx + day_col_w, hdr_y + hdr_h)],
                  fill=GRID_LINE, width=1)

    # Hour rows — Y is tracked cumulatively so each row uses its own height.
    body_y = hdr_y + hdr_h
    ry = body_y
    rows_drawn = 0
    for ri, hh in enumerate(hours):
        rh = row_heights[hh]
        if ry + rh > grid_max_y:
            break
        rows_drawn += 1
        # Alternating row background
        draw.rectangle([grid_x, ry, grid_x + grid_w, ry + rh],
                       fill=ROW_ALT if ri % 2 else ROW_BASE)
        # Time-label cell
        draw.rectangle([grid_x, ry, grid_x + time_col_w, ry + rh], fill=TIME_BG)
        time_lbl = f"{hh:02d}:00"
        tw = draw.textlength(time_lbl, font=F_TIME)
        draw.text((grid_x + (time_col_w - tw)//2, ry + (rh - 22)//2),
                  time_lbl, font=F_TIME, fill=TIME_FG)
        draw.line([(grid_x + time_col_w, ry),
                   (grid_x + time_col_w, ry + rh)],
                  fill=GRID_LINE, width=1)

        # Per-day cells: stack the PKT windows that start in this hour.
        # 1-px inset on the left/right keeps the rendered text from touching
        # the column divider when its glyph is at the cell's edge.
        for di in range(n_cols):
            cx = grid_x + time_col_w + di * day_col_w
            wins_in = day_hour_windows[di].get(hh, []) if di < len(days) else []
            # Cap to what actually fits this row's height (row may have been
            # scaled down to fit the page) so window text NEVER overflows into
            # the neighbouring row. If more windows exist than fit, the last
            # visible slot shows a "+N more" summary instead of colliding text.
            fit = max(1, (rh - 8) // 19)
            if len(wins_in) > fit:
                shown = wins_in[:max(0, fit - 1)]
                overflow = len(wins_in) - len(shown)
            else:
                shown = wins_in
                overflow = 0
            for wi, w in enumerate(shown):
                tw = draw.textlength(w, font=F_CELL)
                pos_x = cx + max(2, (day_col_w - int(tw)) // 2)
                pos = (pos_x, ry + 6 + wi * 19)
                draw.text(pos, w, font=F_CELL, fill=GREEN_TXT)
                sharp_queue.append((pos, w))
            if overflow:
                msg = f"+{overflow} more"
                tw = draw.textlength(msg, font=F_CELL)
                draw.text((cx + max(2, (day_col_w - int(tw)) // 2),
                           ry + 6 + len(shown) * 19),
                          msg, font=F_CELL, fill=SUB_GREY)
            draw.line([(cx + day_col_w, ry),
                       (cx + day_col_w, ry + rh)],
                      fill=GRID_LINE, width=1)
        # Horizontal separator
        draw.line([(grid_x, ry + rh),
                   (grid_x + grid_w, ry + rh)], fill=GRID_LINE, width=1)
        ry += rh

    # Outer border of the table
    draw.rectangle([grid_x, hdr_y, grid_x + grid_w, ry],
                   outline=cfg["border_rgb"], width=2)

    # ── Neon-glow composite pass (green cell text only) ─────────────────────
    # Tight halo radii keep the glow inside each column so adjacent cells
    # don't visually bleed into each other.
    glow_inner = glow.filter(ImageFilter.GaussianBlur(radius=1.0))
    glow_outer = glow.filter(ImageFilter.GaussianBlur(radius=2.5))
    img = Image.alpha_composite(img, glow_outer)
    img = Image.alpha_composite(img, glow_inner)
    draw = ImageDraw.Draw(img)
    for pos, w in sharp_queue:
        draw.text(pos, w, font=F_CELL, fill=GREEN_TXT)

    # ── Footer strip ────────────────────────────────────────────────────────
    fy = H - 62
    draw.rectangle([0, fy, W, H], fill=BAR_BG)

    # NOTES (left)
    draw.text((12, fy + 4), "NOTES:", font=F_FOOT_HDR, fill=TITLE)
    draw.text((12, fy + 20), "All times in Pakistan Standard Time (PKT, UTC+5)",
              font=F_FOOT, fill=SUB_GREY)
    draw.text((12, fy + 34), "Green = Satellite Above (Coverage)",
              font=F_FOOT, fill=GREEN_TXT)
    draw.text((12, fy + 48), "Black/Dark = Blind Window (No Coverage)",
              font=F_FOOT, fill=SUB_GREY)

    # Info circle (centre)
    info_x = 420
    draw.ellipse([info_x, fy + 18, info_x + 22, fy + 40], outline=CYAN, width=2)
    iw = draw.textlength("i", font=F_FOOT_HDR)
    draw.text((info_x + (22 - iw)//2, fy + 20), "i", font=F_FOOT_HDR, fill=CYAN)
    draw.text((info_x + 32, fy + 10), "Optical satellites capture only in daylight",
              font=F_FOOT, fill=SUB_GREY)
    draw.text((info_x + 32, fy + 24), "SAR satellites can capture day & night",
              font=F_FOOT, fill=SUB_GREY)
    draw.text((info_x + 32, fy + 38), "Schedule generated using your 15-day forecast data",
              font=F_FOOT, fill=SUB_GREY)

    # Total coverage (right-of-centre) — reports the COMBINED view for the
    # current country bucket. Parses both compact time formats produced by
    # _windows(): "HH:MM-MM" (same hour) and "HH:MM-HH:MM" (hour crossing).
    _country_pred = lambda c: _country_match(c, country)
    per_day = []
    for d in days:
        m = 0
        for x in _windows(d, _country_pred):
            try:
                parts = x.split("-")
                if len(parts) != 2:
                    continue
                a, b = parts
                ah, am = map(int, a.split(":"))
                if ":" in b:
                    bh, bm = map(int, b.split(":"))
                else:
                    bh, bm = ah, int(b)
                m += (bh * 60 + bm) - (ah * 60 + am)
            except Exception:
                pass
        per_day.append(m)
    avg_cov   = sum(per_day) / max(1, len(per_day))
    avg_blind = max(0, 24 * 60 - avg_cov)

    tc_x = 1000
    draw.rectangle([tc_x, fy + 6, tc_x + 280, fy + 56],
                   outline=(170, 178, 190), width=1)
    hdr = f"TOTAL COVERAGE ({country_label})"
    hw = draw.textlength(hdr, font=F_FOOT_HDR)
    draw.text((tc_x + (280 - hw)//2, fy + 9), hdr, font=F_FOOT_HDR, fill=CYAN)
    draw.text((tc_x + 12, fy + 26),
              f"Avg. coverage per day:    ~ {avg_cov:.0f} minutes",
              font=F_FOOT, fill=TITLE)
    draw.text((tc_x + 12, fy + 40),
              f"Avg. blind window per day:  ~ {int(avg_blind)//60} h {int(avg_blind)%60} m",
              font=F_FOOT, fill=TITLE)

    # Dish glyph + Generated on (far right)
    gx = 1320
    draw.arc([gx, fy + 14, gx + 32, fy + 46], 180, 360,
             fill=(120, 128, 140), width=2)
    draw.line([gx + 16, fy + 30, gx + 16, fy + 52],
              fill=(120, 128, 140), width=2)
    draw.text((gx + 42, fy + 18), "Generated on:", font=F_FOOT, fill=SUB_GREY)
    when = (_dt.datetime.now(_dt.timezone.utc) + _PKT).strftime("%d %b %Y %H:%M PKT")
    draw.text((gx + 42, fy + 33), when, font=F_FOOT_HDR, fill=TITLE)

    img.convert("RGB").save(out_png, optimize=True,
                            dpi=(1600 / 14.4, 1600 / 14.4))
    return out_png


def render_word(forecast: dict | None = None,
                output_path: str | None = None) -> str | None:
    if forecast is None:
        if not os.path.isfile(FORECAST_JSON):
            print("[FORECAST-DOC] no forecast JSON found — run build_forecast() first")
            return None
        with open(FORECAST_JSON, "r", encoding="utf-8") as f:
            forecast = json.load(f)

    try:
        from docx import Document
        from docx.shared import RGBColor, Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("[FORECAST-DOC] python-docx not installed — skipping Word output")
        return None

    RED       = RGBColor(0xC8, 0x10, 0x2E)
    DARK_GREY = RGBColor(0x33, 0x33, 0x33)
    BLUE      = RGBColor(0x1F, 0x3A, 0x5F)

    doc = Document()
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.0)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)

    title = doc.add_heading(
        "ATLAS — Pakistan 15-Day Hi-Res EO Surveillance Forecast", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in title.runs:
        r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    gen_pkt = _pkt_datetime(forecast['generated_at']) + " PKT"
    sub.add_run(
        f"Window: {forecast['start_date']}  →  {forecast['end_date']}     "
        f"|   Generated: {gen_pkt}"
    ).font.size = Pt(10)

    # Timezone declaration — every time in this document is PKT (UTC + 5h).
    tz = doc.add_paragraph()
    tz.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tzr = tz.add_run(
        "All times in this document are Pakistan Standard Time (PKT, UTC+5h, no DST)."
    )
    tzr.italic = True
    tzr.font.size = Pt(9)
    tzr.font.color.rgb = BLUE

    # ── Executive summary ───────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    total_cross    = sum(d["totals"]["crossings"]            for d in forecast["days"])
    total_overhead = sum(d["totals"]["overhead_crossings"]   for d in forecast["days"])
    total_tilt     = sum(d["totals"]["tilt_range_crossings"] for d in forecast["days"])
    total_indian   = sum(d["totals"]["indian_crossings"]     for d in forecast["days"])
    ind_overhead   = sum(d["totals"]["indian_overhead"]      for d in forecast["days"])
    ind_tilt       = sum(d["totals"]["indian_tilt_range"]    for d in forecast["days"])
    total_israeli  = sum(d["totals"].get("israeli_crossings", 0)   for d in forecast["days"])
    isr_overhead   = sum(d["totals"].get("israeli_overhead", 0)    for d in forecast["days"])
    isr_tilt       = sum(d["totals"].get("israeli_tilt_range", 0)  for d in forecast["days"])
    total_blind    = sum(d["totals"]["blind_minutes"]        for d in forecast["days"])
    longest_blind  = max((d["totals"]["longest_blind_min"] for d in forecast["days"]),
                         default=0)

    # Israeli sat count from catalog (run-time, not stored in forecast JSON
    # — keeps backward compatibility with pre-Israel-split forecasts).
    try:
        from hires_eo_satellites import HIRES_EO_SATELLITES as _HES
        israeli_sat_count = sum(1 for s in _HES if s.country == "Israel")
    except Exception:
        israeli_sat_count = 0

    summary_rows = [
        ("Forecast window",            f"{forecast['start_date']} -> {forecast['end_date']} ({len(forecast['days'])} days)"),
        ("Surveillance region",        "Pakistan box 23.5°N-37.5°N, 60.5°E-77.5°E"),
        ("Tilt-range model",           "PER-SATELLITE standoff = altitude x tan(max_tilt) — see Imaging Capability table below"),
        ("Tracked HiRes EO satellites", str(forecast["satellite_count"])),
        ("Indian (ISRO) satellites",   str(forecast["indian_satellite_count"])),
        ("Israeli satellites",         f"{israeli_sat_count}  (Ofeq mil-recon series + EROS-C3 commercial)"),
        ("Total predicted passes",     f"{total_cross}  ({total_overhead} overhead + {total_tilt} tilt-range)"),
        ("Indian passes",              f"{total_indian}  ({ind_overhead} overhead + {ind_tilt} tilt-range)"),
        ("Israeli passes",             f"{total_israeli}  ({isr_overhead} overhead + {isr_tilt} tilt-range)"),
        ("Total blind time",           f"{total_blind:,.0f} min  ({total_blind/60:.1f} hours)"),
        ("Longest single blind window", f"{longest_blind:.0f} minutes"),
        ("Propagator",                 forecast["propagator"]),
    ]
    tbl = doc.add_table(rows=len(summary_rows), cols=2)
    tbl.style = "Light Grid Accent 1"
    for i, (k, v) in enumerate(summary_rows):
        tbl.cell(i, 0).text = k
        tbl.cell(i, 1).text = str(v)
        tbl.cell(i, 0).paragraphs[0].runs[0].bold = True
        if "Indian" in k:
            for c in tbl.row_cells(i):
                for run in c.paragraphs[0].runs:
                    run.font.color.rgb = RED
                    run.bold = True

    # ── Strategic Site Coverage chart (FEAT-004) ────────────────────────────
    # For every catalogued site, count how many ≤1m passes target it across
    # the 15-day window, broken down by SAR / Optical and overhead / tilt.
    # Operator scans this section first: "which assets are most exposed".
    doc.add_heading("Strategic Site Coverage — Pakistan (15-day window)", level=1)
    cov_intro = doc.add_paragraph()
    cov_intro.add_run(
        "Each row is a publicly-known strategic facility. The count columns "
        "show how many ≤1 m optical / SAR / military satellite passes can "
        "image the site over the 15-day forecast window — a pass 'targets' "
        "a site iff the satellite's ground sub-point comes within that "
        "specific satellite's tilt-standoff radius of the site at some "
        "moment during the pass arc. Sub-classes:  OVR = direct overhead, "
        "TLT = off-nadir / side-look reach. Source coordinates are "
        "open-source only (Wikipedia / public references)."
    ).font.size = Pt(9)
    cov_intro.runs[0].italic = True

    # Build per-site aggregate counts from every crossing in the forecast.
    site_stats: dict[str, dict] = {}
    for d in forecast["days"]:
        for c in d.get("crossings", []):
            sensor_is_sar = "SAR" in (c.get("sensor", "") or "").upper()
            ptyp_over = c.get("pass_type") == "overhead"
            for t in c.get("targeted_sites") or []:
                key = t["name"]
                entry = site_stats.setdefault(key, {
                    "name": t["name"], "city": t["city"], "tier": t["tier"],
                    "category": t["category"],
                    "opt_over": 0, "opt_tilt": 0,
                    "sar_over": 0, "sar_tilt": 0,
                    "total": 0,
                    "min_dist_km": 1e9,
                })
                if sensor_is_sar:
                    entry["sar_over" if ptyp_over else "sar_tilt"] += 1
                else:
                    entry["opt_over" if ptyp_over else "opt_tilt"] += 1
                entry["total"] += 1
                if t["min_dist_km"] < entry["min_dist_km"]:
                    entry["min_dist_km"] = t["min_dist_km"]

    # Render the catalog in tier order — tier-1 first (most relevant), each
    # sorted by total pass count descending. Always include every catalog
    # site even if zero passes — operator wants to see "nothing imaged Khushab".
    try:
        from pakistan_strategic_sites import STRATEGIC_SITES as _ALL_SITES
    except ImportError:
        _ALL_SITES = []

    cov_cols = ["Site", "City", "Tier", "Category",
                "Opt OVR", "Opt TLT", "SAR OVR", "SAR TLT",
                "Total", "Closest (km)"]
    cov_tbl = doc.add_table(rows=1 + len(_ALL_SITES), cols=len(cov_cols))
    cov_tbl.style = "Light Grid Accent 1"
    for i, h in enumerate(cov_cols):
        cell = cov_tbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    # Sort: tier 1 first, then total passes descending, then name
    ordered = sorted(_ALL_SITES, key=lambda s: (
        s["tier"], -site_stats.get(s["name"], {}).get("total", 0), s["name"]))
    TARGET_RED = RGBColor(0xC8, 0x10, 0x2E)
    for i, s in enumerate(ordered, 1):
        st = site_stats.get(s["name"], {})
        row = cov_tbl.rows[i].cells
        row[0].text = s["name"][:42]
        row[1].text = s["city"]
        row[2].text = f"T{s['tier']}"
        row[3].text = s["category"]
        row[4].text = str(st.get("opt_over", 0))
        row[5].text = str(st.get("opt_tilt", 0))
        row[6].text = str(st.get("sar_over", 0))
        row[7].text = str(st.get("sar_tilt", 0))
        row[8].text = str(st.get("total", 0))
        cd = st.get("min_dist_km", None)
        row[9].text = f"{cd:.0f}" if cd is not None and cd < 1e8 else "—"
        for cell in row:
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8)
        # Highlight tier-1 rows in red + bold-red the total if it's > 0
        if s["tier"] == 1:
            for r in row[0].paragraphs[0].runs:
                r.font.color.rgb = TARGET_RED
                r.bold = True
        if st.get("total", 0) >= 10:
            for r in row[8].paragraphs[0].runs:
                r.font.color.rgb = TARGET_RED
                r.bold = True

    # Headline summary line — total tier-1 exposures across all 15 days
    t1_total = sum(st["total"] for k, st in site_stats.items()
                   if st.get("tier") == 1)
    t2_total = sum(st["total"] for k, st in site_stats.items()
                   if st.get("tier") == 2)
    headline = doc.add_paragraph()
    hr = headline.add_run(
        f"Headline: {t1_total} tier-1 site-targeting passes and "
        f"{t2_total} tier-2 site-targeting passes over the 15-day window. "
        f"Cells in red indicate ≥10 reachable passes — high-exposure assets."
    )
    hr.italic = True
    hr.font.size = Pt(9)
    hr.font.color.rgb = TARGET_RED

    # ── Per-satellite Imaging Capability (tilt + standoff, ≤1m only) ────────
    # The table the user explicitly asked for: every sub-metre optical/SAR
    # satellite this forecast tracks, with its nominal altitude, maximum
    # off-nadir / incidence angle, and the resulting standoff radius.
    doc.add_heading("Imaging Capability — Tilt & Standoff per Satellite", level=1)
    cap_intro = doc.add_paragraph()
    cap_intro.add_run(
        "Each satellite is listed with its operational maximum off-nadir "
        "(optical) or maximum incidence (SAR) angle, plus the resulting "
        "ground standoff radius: altitude x tan(max_tilt). Only satellites "
        "with ≤1 m resolution are shown — these are the only platforms whose "
        "imagery is operationally consequential. SAR/optical only."
    ).font.size = Pt(9)
    cap_intro.runs[0].italic = True

    try:
        from hires_eo_satellites import HIRES_EO_SATELLITES, TILT_SPECS, standoff_km as _so
    except ImportError:
        HIRES_EO_SATELLITES, TILT_SPECS = [], {}
        _so = lambda _n: None
    sub1m = [s for s in HIRES_EO_SATELLITES
             if s.resolution_m <= 1.0 and s.sensor_category in ("Optical", "SAR", "Military")]
    # Split by sensor type — Optical above / SAR below per operator request.
    sub1m_optical = [s for s in sub1m if "SAR" not in s.sensor.upper()]
    sub1m_sar     = [s for s in sub1m if "SAR" in s.sensor.upper()]
    # Sort each bucket by standoff descending
    sub1m_optical.sort(key=lambda s: (_so(s.norad_id) or 0), reverse=True)
    sub1m_sar.sort    (key=lambda s: (_so(s.norad_id) or 0), reverse=True)

    cap_cols = ["Satellite", "Country", "Operator", "Sensor", "Res (m)",
                "Alt (km)", "Max Tilt (°)", "Standoff (km)"]

    def _render_capability_block(label: str, sats: list):
        if not sats:
            doc.add_paragraph(f"  (no {label.lower()} satellites in catalog)")
            return
        sh = doc.add_heading(f"{label}  ({len(sats)} satellites)", level=2)
        cap_tbl = doc.add_table(rows=1 + len(sats), cols=len(cap_cols))
        cap_tbl.style = "Light Grid Accent 1"
        for i, h in enumerate(cap_cols):
            cell = cap_tbl.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        for i, s in enumerate(sats, 1):
            spec = TILT_SPECS.get(str(s.norad_id), {})
            row = cap_tbl.rows[i].cells
            row[0].text = s.name
            row[1].text = s.country
            row[2].text = s.operator
            row[3].text = s.sensor_category
            row[4].text = f"{s.resolution_m:.2f}"
            row[5].text = str(spec.get("altitude_km", "-"))
            row[6].text = str(spec.get("max_tilt_deg", "-"))
            so = _so(s.norad_id)
            row[7].text = f"{so:.0f}" if so is not None else "-"
            for cell in row:
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(8)
            # Indian rows in red
            if s.country == "India":
                for cell in (row[0], row[1]):
                    for r in cell.paragraphs[0].runs:
                        r.font.color.rgb = RED
                        r.bold = True
            # Israeli rows in blue (distinct from India's red, so the eye can
            # tell the two regional threat-actor blocs apart at a glance)
            elif s.country == "Israel":
                _ISR_BLUE = RGBColor(0x1F, 0x60, 0xC0)
                for cell in (row[0], row[1]):
                    for r in cell.paragraphs[0].runs:
                        r.font.color.rgb = _ISR_BLUE
                        r.bold = True

    # OPTICAL above
    _render_capability_block("Optical", sub1m_optical)
    # SAR below
    _render_capability_block("SAR", sub1m_sar)

    # ── 15-Day Blind-Window Timetables — INDIAN vs OTHER-COUNTRY ───────────
    # Per the boss's instruction: show ONLY the exact blind time — the hours
    # when NO satellite of a given group has Pakistan in its imaging
    # footprint — split into two day x hour grids: Indian satellites, and
    # every other country. Covered hours are left BLANK so the eye sees
    # nothing but the blind windows.
    try:
        from docx.enum.section import WD_ORIENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        _HAVE_GRID = True
    except ImportError:
        _HAVE_GRID = False

    def _shade(cell, hex_fill):
        if not _HAVE_GRID:
            return
        tcPr = cell._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), hex_fill)
        tcPr.append(shd)

    def _cell_style(m: float):
        """minutes-blind -> (fill_hex, cell_text, white_text). Covered hours
        are left BLANK so the grid shows ONLY the blind time."""
        if m <= 0.5:
            return ("FFFFFF", "", False)           # covered — blank
        if m < 20:
            return ("F5CBA7", f"{m:.0f}", False)   # light blind — amber
        if m < 45:
            return ("E59866", f"{m:.0f}", False)   # partial blind
        if m < 59.5:
            return ("CB4335", f"{m:.0f}", True)    # mostly blind
        return ("A93226", "60", True)              # fully blind hour — deep red

    from datetime import datetime as _dtg, timedelta as _tdg

    def _blind_timetable(win_key: str, heading: str, intro_text: str):
        """Render ONE day x hour blind-time grid from the blind windows under
        `win_key` (e.g. 'india_blind_windows'). Rows = PKT dates, columns =
        the 24 hours. Covered hours stay blank — only blind minutes show."""
        doc.add_heading(heading, level=1)
        _intro = doc.add_paragraph()
        _intro.add_run(intro_text).font.size = Pt(9)
        _intro.runs[0].italic = True

        # PKT day x hour grid: pkt_date -> [24] blind-minute buckets.
        grid: dict[str, list[float]] = {}
        for d in forecast["days"]:
            for bw in d.get(win_key, []):
                try:
                    s = _dtg.strptime(bw["start"], "%Y-%m-%dT%H:%M:%SZ") + _tdg(hours=5)
                    e = _dtg.strptime(bw["end"],   "%Y-%m-%dT%H:%M:%SZ") + _tdg(hours=5)
                except (ValueError, KeyError):
                    continue
                cur = s
                while cur < e:
                    hour_end = cur.replace(minute=0, second=0, microsecond=0) + _tdg(hours=1)
                    seg_end = min(e, hour_end)
                    dkey = cur.strftime("%Y-%m-%d")
                    grid.setdefault(dkey, [0.0] * 24)
                    grid[dkey][cur.hour] += (seg_end - cur).total_seconds() / 60.0
                    cur = seg_end

        if not grid:
            doc.add_paragraph("No blind windows for this satellite group across "
                              "the 15-day window — Pakistan is observed at every "
                              "minute by at least one of its platforms.")
            return

        pkt_dates = sorted(grid)
        tt = doc.add_table(rows=1 + len(pkt_dates), cols=25)
        tt.style = "Table Grid"
        tt.autofit = False
        _date_w = Cm(2.9)
        _hour_w = Cm(1.0)               # 2.9 + 24 x 1.0 = 26.9 cm usable
        if _HAVE_GRID:
            _tblPr = tt._tbl.tblPr
            # Single fixed-layout element so Word honours the column widths
            # and the grid fills the full landscape page width.
            for _el in _tblPr.findall(qn("w:tblLayout")):
                _tblPr.remove(_el)
            _layout = OxmlElement("w:tblLayout")
            _layout.set(qn("w:type"), "fixed")
            _tblPr.append(_layout)
        # Header row — PKT Date + hours 00..23
        hdr = tt.rows[0].cells
        hdr[0].text = "PKT Date"
        for h in range(24):
            hdr[h + 1].text = f"{h:02d}"
        for c in hdr:
            _shade(c, "1F3A5F")
            c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in c.paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(6)
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        # Data rows — one per PKT date
        for ri, dkey in enumerate(pkt_dates, 1):
            cells = tt.rows[ri].cells
            cells[0].text = dkey
            cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cells[0].paragraphs[0].runs:
                r.bold = True
                r.font.size = Pt(6)
            for h in range(24):
                m = grid[dkey][h]
                fill, txt, white = _cell_style(m)
                cell = cells[h + 1]
                cell.text = txt
                _shade(cell, fill)
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in cell.paragraphs[0].runs:
                    r.font.size = Pt(6)
                    if white:
                        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                        r.bold = True
        # Fixed column widths — set on every cell AND every column.
        for col_idx, w in [(0, _date_w)] + [(h, _hour_w) for h in range(1, 25)]:
            try:
                tt.columns[col_idx].width = w
            except (IndexError, AttributeError):
                pass
        for row in tt.rows:
            row.cells[0].width = _date_w
            for h in range(1, 25):
                row.cells[h].width = _hour_w
        # Taller rows so the grid fills the landscape page like a wall planner.
        try:
            from docx.enum.table import WD_ROW_HEIGHT_RULE
            tt.rows[0].height = Cm(0.6)
            tt.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            for row in tt.rows[1:]:
                row.height = Cm(0.85)
                row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
        except ImportError:
            pass
        # Legend
        leg = doc.add_paragraph()
        leg.add_run("Legend:  ").bold = True
        leg.add_run("blank = covered (a satellite of this group is above "
                    "Pakistan)      1-19 / 20-44 / 45-59 = blind minutes in "
                    "that PKT hour      60 = fully blind hour.      "
                    "Cell value = exact blind minutes within that PKT hour.")
        for r in leg.runs:
            r.font.size = Pt(8)

    # ── 15-day Indian satellite coverage timetable ──────────────────────────
    # Rendered as a pixel-exact PNG (wall-planner format) and embedded as a
    # single full-bleed landscape image. Replaces the older 24-hour blind grids.
    if _HAVE_GRID:
        land = doc.add_section()
        land.orientation = WD_ORIENT.LANDSCAPE
        land.page_width  = Cm(29.7)
        land.page_height = Cm(21.0)
        land.left_margin = land.right_margin = Cm(0.5)
        land.top_margin  = land.bottom_margin = Cm(0.5)
    else:
        doc.add_page_break()

    # One landscape page per category (Combined / Optical / SAR). Each PNG
    # is also written to documents/ so the operator can view / print them
    # standalone without opening Word.
    _cat_files = [
        ("combined", "Pakistan_15Day_Coverage_Combined.png"),
        ("optical",  "Pakistan_15Day_Coverage_Optical.png"),
        ("sar",      "Pakistan_15Day_Coverage_SAR.png"),
    ]
    _png_paths_for_pdf: list[str] = []
    for _idx, (_cat, _fname) in enumerate(_cat_files):
        _png_path = os.path.join(DOCS_DIR, _fname)
        try:
            _render_category_png(forecast, _cat, _png_path)
            doc.add_picture(_png_path, width=Cm(29.0))
            _png_paths_for_pdf.append(_png_path)
        except Exception as _e:
            _err = doc.add_paragraph()
            _er = _err.add_run(
                f"[Coverage timetable image ({_cat}) could not be rendered: {_e}]"
            )
            _er.italic = True
            _er.font.color.rgb = RGBColor(0xC8, 0x10, 0x2E)
        # Page break between category pages (but not after the last one —
        # the portrait section break below already starts a new page).
        if _idx < len(_cat_files) - 1:
            doc.add_page_break()

    # ── Plotter-ready multipage PDF ────────────────────────────────────────
    # Bundle the 3 category PNGs into a single landscape PDF (3 pages).
    # resolution=72 dpi puts each page at ~22"×14" — roughly A2 landscape —
    # which a plotter can scale up to A1 or A0 with crisp rendering.
    try:
        from PIL import Image as _Img
        if _png_paths_for_pdf:
            _imgs = []
            for p in _png_paths_for_pdf:
                _s = _Img.open(p); _imgs.append(_s.convert("RGB")); _s.close()
            _pdf_path = os.path.join(DOCS_DIR,
                                     "Pakistan_15Day_Coverage_Timetable.pdf")
            _imgs[0].save(_pdf_path, format="PDF", resolution=72.0,
                          save_all=True, append_images=_imgs[1:])
            print(f"[FORECAST-DOC] wrote {_pdf_path}")
            _cleanup_pngs(_png_paths_for_pdf)   # embedded in the PDF — drop PNGs
    except Exception as _pe:
        print(f"[FORECAST-DOC] PDF bundle skipped: {_pe}")

    # Return to portrait for the remaining sections.
    if _HAVE_GRID:
        port = doc.add_section()
        port.orientation = WD_ORIENT.PORTRAIT
        port.page_width  = Cm(21.0)
        port.page_height = Cm(29.7)
        port.left_margin = port.right_margin = Cm(2.0)
        port.top_margin  = port.bottom_margin = Cm(2.5)

    # Per-day blind-time roll-up — one row per day, blind minutes only,
    # split Indian / Israeli / other-country / combined.
    doc.add_heading("Daily Blind-Time Roll-up", level=2)
    rollup_cols = ["Date", "Indian-sat blind (min)", "Israeli-sat blind (min)",
                   "Other-country blind (min)", "Combined blind (min)",
                   "Longest blind gap (min)"]
    rollup_rows = []
    for d in forecast["days"]:
        t = d["totals"]
        rollup_rows.append((
            d["date"],
            t.get("india_blind_minutes", 0),
            t.get("israel_blind_minutes", 0),
            t.get("other_blind_minutes", 0),
            t.get("blind_minutes", 0),
            t.get("longest_blind_min", 0),
        ))
    rtbl = doc.add_table(rows=1 + len(rollup_rows), cols=len(rollup_cols))
    rtbl.style = "Light Grid Accent 1"
    for i, h in enumerate(rollup_cols):
        cell = rtbl.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.size = Pt(9)
    _max_ind = max((r[1] for r in rollup_rows), default=0)
    _max_isr = max((r[2] for r in rollup_rows), default=0)
    _max_oth = max((r[3] for r in rollup_rows), default=0)
    for i, (date, ind_b, isr_b, oth_b, comb_b, longest) in enumerate(rollup_rows, 1):
        row = rtbl.rows[i].cells
        row[0].text = date
        row[1].text = f"{ind_b:.0f}"
        row[2].text = f"{isr_b:.0f}"
        row[3].text = f"{oth_b:.0f}"
        row[4].text = f"{comb_b:.0f}"
        row[5].text = f"{longest:.0f}"
        for cell in row:
            for r in cell.paragraphs[0].runs:
                r.font.size = Pt(8)
        # Red the worst (longest blind) day in each country column.
        if ind_b > 0 and ind_b == _max_ind:
            for r in row[1].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); r.bold = True
        if isr_b > 0 and isr_b == _max_isr:
            for r in row[2].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); r.bold = True
        if oth_b > 0 and oth_b == _max_oth:
            for r in row[3].paragraphs[0].runs:
                r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B); r.bold = True

    # ── Daily breakdown ─────────────────────────────────────────────────────
    for day_idx, day in enumerate(forecast["days"]):
        doc.add_page_break()
        h = doc.add_heading(f"{day['date']}", level=1)
        for r in h.runs:
            r.font.color.rgb = BLUE

        # Provenance line — shows the user exactly which run produced this
        # prediction. day 0 = "today" (using freshest TLEs); offsets > 0 are
        # progressively coarser SGP4 propagation.
        prov = doc.add_paragraph()
        pr = prov.add_run(
            f"Predicted on {forecast['start_date']} using CelesTrak TLEs "
            f"current at {_pkt_datetime(forecast['generated_at'])} PKT — "
            + ("freshest TLE epoch (day-0 — most accurate)" if day_idx == 0
               else f"propagated {day_idx} day(s) ahead (SGP4 error grows ~linearly)")
        )
        pr.italic = True
        pr.font.color.rgb = DARK_GREY
        pr.font.size = Pt(9)

        t = day["totals"]
        doc.add_paragraph(
            f"Passes: {t['crossings']} total "
            f"({t['overhead_crossings']} overhead + {t['tilt_range_crossings']} tilt-range)    "
            f"Unique satellites: {t['unique_sats']}    "
            f"Blind: {t['blind_minutes']:.0f} min    "
            f"Longest blind: {t['longest_blind_min']:.0f} min"
        )

        if not day["crossings"]:
            p = doc.add_paragraph()
            r = p.add_run("WARNING: No HiRes EO daytime passes predicted for this day "
                          "- Pakistan is fully blind to <=1m optical observation.")
            r.bold = True
            r.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
            continue

        # ── INDIAN section (red) — split into OPTICAL above / SAR below ──
        india_list = day.get("india_crossings", [])
        india_opt, india_sar = _split_optical_sar(india_list)
        ih = doc.add_heading(
            f"INDIA (ISRO) - {t['indian_crossings']} passes "
            f"({len(india_opt)} optical + {len(india_sar)} SAR  ·  "
            f"{t['indian_overhead']} overhead + {t['indian_tilt_range']} tilt-range)",
            level=2)
        for r in ih.runs:
            r.font.color.rgb = RED

        if india_list:
            # OPTICAL — above
            doc.add_heading(f"  Optical ({len(india_opt)} passes)", level=3)
            if india_opt:
                _add_crossings_table(doc, india_opt, RED, Pt, RGBColor, indian=True)
            else:
                doc.add_paragraph("    (no Indian optical passes this day)")
            # SAR — below
            doc.add_heading(f"  SAR ({len(india_sar)} passes)", level=3)
            if india_sar:
                _add_crossings_table(doc, india_sar, RED, Pt, RGBColor, indian=True)
            else:
                doc.add_paragraph("    (no Indian SAR passes this day)")
        else:
            doc.add_paragraph("  (no Indian passes this day)")

        # ── ISRAELI section (blue) — same Optical/SAR split. Israel gets
        # its own block (separate from OTHER COUNTRIES) because it operates
        # a dedicated mil-recon fleet (Ofeq series + EROS commercial) and
        # OSINT analysts track its passes independently.
        _ISR_BLUE = RGBColor(0x1F, 0x60, 0xC0)
        israel_list = day.get("israel_crossings", [])
        isr_opt, isr_sar = _split_optical_sar(israel_list)
        isr_count_total = t.get("israeli_crossings", len(israel_list))
        isr_oh   = t.get("israeli_overhead", 0)
        isr_tilt = t.get("israeli_tilt_range", 0)
        ish = doc.add_heading(
            f"ISRAEL (Ofeq / EROS) - {isr_count_total} passes "
            f"({len(isr_opt)} optical + {len(isr_sar)} SAR  ·  "
            f"{isr_oh} overhead + {isr_tilt} tilt-range)",
            level=2)
        for r in ish.runs:
            r.font.color.rgb = _ISR_BLUE

        if israel_list:
            doc.add_heading(f"  Optical ({len(isr_opt)} passes)", level=3)
            if isr_opt:
                _add_crossings_table(doc, isr_opt, _ISR_BLUE, Pt, RGBColor, indian=False)
            else:
                doc.add_paragraph("    (no Israeli optical passes this day)")
            doc.add_heading(f"  SAR ({len(isr_sar)} passes)", level=3)
            if isr_sar:
                _add_crossings_table(doc, isr_sar, _ISR_BLUE, Pt, RGBColor, indian=False)
            else:
                doc.add_paragraph("    (no Israeli SAR passes this day)")
        else:
            doc.add_paragraph("  (no Israeli passes this day)")

        # ── OTHER countries section — same Optical/SAR split (excludes
        # India and Israel — those have their own blocks above).
        other_list = day.get("other_crossings", [])
        other_opt, other_sar = _split_optical_sar(other_list)
        other_count_total = (t['crossings']
                             - t['indian_crossings']
                             - t.get('israeli_crossings', 0))
        oh = doc.add_heading(
            f"OTHER COUNTRIES - {other_count_total} passes "
            f"({len(other_opt)} optical + {len(other_sar)} SAR  ·  "
            f"{t['other_overhead']} overhead + {t['other_tilt_range']} tilt-range)",
            level=2)
        for r in oh.runs:
            r.font.color.rgb = BLUE

        if other_list:
            doc.add_heading(f"  Optical ({len(other_opt)} passes)", level=3)
            if other_opt:
                _add_crossings_table(doc, other_opt, BLUE, Pt, RGBColor, indian=False)
            else:
                doc.add_paragraph("    (no other-country optical passes this day)")
            doc.add_heading(f"  SAR ({len(other_sar)} passes)", level=3)
            if other_sar:
                _add_crossings_table(doc, other_sar, BLUE, Pt, RGBColor, indian=False)
            else:
                doc.add_paragraph("    (no other-country SAR passes this day)")
        else:
            doc.add_paragraph("  (no other-country passes this day)")

        # ── Blind windows section — combined AND per-country breakdown ──
        # Combined = Pakistan invisible to literally every satellite
        #            (operationally: nobody is watching at all).
        # Per-country = Pakistan invisible to that country's fleet specifically
        #               (operationally: that country has no eye on Pakistan
        #               during this window, even if other countries do).
        doc.add_heading("Observation gap (blind) windows", level=2)

        _ISR_BLUE = RGBColor(0x1F, 0x60, 0xC0)

        def _render_blind_table(label: str, windows: list,
                                color: "RGBColor | None",
                                empty_msg: str):
            """Render one blind-window table under a sub-heading. Times are
            already UTC strings in the JSON — convert to PKT for display."""
            sh = doc.add_heading(label, level=3)
            if color is not None:
                for r in sh.runs:
                    r.font.color.rgb = color
            if not windows:
                doc.add_paragraph(empty_msg)
                return
            bcols = ["Start (PKT)", "End (PKT)", "Duration (min)"]
            btbl = doc.add_table(rows=1 + len(windows), cols=len(bcols))
            btbl.style = "Light Grid Accent 1"
            for i, h in enumerate(bcols):
                cell = btbl.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            for i, b in enumerate(windows, 1):
                row = btbl.rows[i].cells
                row[0].text = _pkt_time(b["start"])
                row[1].text = _pkt_time(b["end"])
                row[2].text = f"{b['duration_min']:.0f}"
                for cell in row:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)

        # 1. Combined (all satellites) — when NOBODY is watching
        _render_blind_table(
            "Combined — Pakistan invisible to ALL satellites",
            day.get("blind_windows", []),
            color=None,
            empty_msg="None — Pakistan is observed at every minute of the day."
        )

        # 2. Indian-only — when India's fleet is blind (even if other countries see us)
        _render_blind_table(
            "India only — Pakistan invisible to ISRO fleet specifically",
            day.get("india_blind_windows", []),
            color=RED,
            empty_msg="None — Indian fleet covers Pakistan continuously today."
        )

        # 3. Israeli-only — when Israel's fleet is blind
        _render_blind_table(
            "Israel only — Pakistan invisible to Ofeq / EROS fleet specifically",
            day.get("israel_blind_windows", []),
            color=_ISR_BLUE,
            empty_msg="None — Israeli fleet covers Pakistan continuously today."
        )

        # 4. Other-country (excludes India + Israel) — when only non-India/Israel are blind
        _render_blind_table(
            "Other countries only — Pakistan invisible to non-India/Israel sats",
            day.get("other_blind_windows", []),
            color=BLUE,
            empty_msg="None — other-country fleet covers Pakistan continuously today."
        )

    # ── PREDICTION ACCURACY (track record) ──────────────────────────────────
    try:
        import accuracy_tracker
        acc = accuracy_tracker.load_accuracy_record()
    except Exception:
        acc = None
    if acc and acc.get("by_lookback"):
        doc.add_page_break()
        h = doc.add_heading("Prediction Accuracy — Track Record", level=1)
        for r in h.runs:
            r.font.color.rgb = BLUE
        doc.add_paragraph(
            "How well past predictions matched what was later observed. "
            "'Actual' for any past day is taken from the freshest archive made "
            "on that day itself (offset 0 of that day's forecast, propagated "
            "from a TLE epoch within hours of the target date — our most "
            "trustworthy single estimate).",
            style="Body Text"
        )
        cols = ["Lookback (days)", "Samples", "Mean accuracy", "Mean entry shift"]
        a_tbl = doc.add_table(rows=1 + len(acc["by_lookback"]), cols=len(cols))
        a_tbl.style = "Light Grid Accent 1"
        for i, c in enumerate(cols):
            cell = a_tbl.rows[0].cells[i]
            cell.text = c
            cell.paragraphs[0].runs[0].bold = True
        for i, b in enumerate(acc["by_lookback"], 1):
            row = a_tbl.rows[i].cells
            row[0].text = f"{b['lookback_days']}"
            row[1].text = str(b["sample_size"])
            row[2].text = f"{b['mean_accuracy_pct']:.1f}%"
            row[3].text = (f"{b['mean_shift_min']:.1f} min"
                           if b.get("mean_shift_min") is not None else "-")
        doc.add_paragraph("")
        # Per-target-date detail (last 30 comparison records for brevity)
        if acc.get("records"):
            doc.add_heading("Recent comparisons (newest first)", level=2)
            recs = sorted(acc["records"],
                          key=lambda r: (r["target_date"], -r["lookback_days"]),
                          reverse=True)[:30]
            cols2 = ["Target date", "Predicted on", "Lookback (d)",
                     "Predicted", "Actual", "Matched", "Accuracy", "Shift (min)"]
            r_tbl = doc.add_table(rows=1 + len(recs), cols=len(cols2))
            r_tbl.style = "Light Grid Accent 1"
            for i, c in enumerate(cols2):
                cell = r_tbl.rows[0].cells[i]
                cell.text = c
                cell.paragraphs[0].runs[0].bold = True
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            for i, r in enumerate(recs, 1):
                row = r_tbl.rows[i].cells
                row[0].text = r["target_date"]
                row[1].text = r["predicted_on"]
                row[2].text = str(r["lookback_days"])
                row[3].text = str(r["predicted_count"])
                row[4].text = str(r["actual_count"])
                row[5].text = str(r["matched"])
                row[6].text = (f"{r['accuracy_pct']:.1f}%"
                               if r.get("accuracy_pct") is not None else "-")
                row[7].text = (f"{r['mean_shift_min']:.1f}"
                               if r.get("mean_shift_min") is not None else "-")
                for cell in row:
                    cell.paragraphs[0].runs[0].font.size = Pt(8)

    # ── Footer ──────────────────────────────────────────────────────────────
    doc.add_paragraph()
    f = doc.add_paragraph()
    f.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = f.add_run(
        "Predictions are computed locally with the sgp4 library from TLEs "
        "fetched on the 'Generated' date above. TLE accuracy degrades with "
        "horizon — entries near day 15 are coarse. Regenerate daily for best "
        "fidelity."
    )
    fr.italic = True
    fr.font.color.rgb = DARK_GREY
    fr.font.size = Pt(9)

    if output_path is None:
        output_path = os.path.join(
            DOCS_DIR,
            f"Pakistan_15Day_Forecast_{forecast['start_date']}.docx"
        )
    doc.save(output_path)
    print(f"[FORECAST-DOC] wrote {output_path}")
    return output_path


def render_report_now() -> str | None:
    """On-demand: render the 15-Day Forecast Word doc (and its Timetable PDF)
    from the current forecast JSON, then refresh the stable *_LATEST.docx
    pointer so "OPEN WORD REPORT" finds it. Returns the LATEST path, or None.

    This is what the GUI "GENERATE REPORT" button calls — report generation is
    no longer a side-effect of the daily run / scan.
    """
    path = render_word()
    if not path:
        return None
    latest = os.path.join(os.path.dirname(path),
                          "Pakistan_15Day_Forecast_LATEST.docx")
    try:
        import shutil
        shutil.copy2(path, latest)
        print(f"[FORECAST-DOC] LATEST -> {os.path.basename(path)}")
    except OSError as e:
        print(f"[FORECAST-DOC] could not refresh LATEST pointer: {e}")
    return latest


# ── CLI ─────────────────────────────────────────────────────────────────────
def run(force_refresh_tles: bool = False, render_doc: bool = True) -> dict:
    """
    Fetch fresh TLEs, build the 15-day forecast, optionally render the Word
    doc. `render_doc=False` is used by run_daily.py so the Word doc can be
    rendered AFTER diff + archive + accuracy_tracker have updated their
    artefacts (so the doc picks up the freshest accuracy data).
    """
    tles = fetch_hires_eo_tles(force_refresh=force_refresh_tles)
    forecast = build_forecast(tles=tles)
    if render_doc:
        render_word(forecast)
    return forecast


if __name__ == "__main__":
    refresh = "--refresh" in sys.argv
    if "--report" in sys.argv:
        # Explicit on-demand report: Word doc + Timetable PDF + LATEST pointer.
        render_report_now()
    else:
        # Default: refresh the forecast JSON only. The Word doc / Timetable PDF
        # are produced on demand (`--report`, or the GUI "GENERATE REPORT"
        # button) — never as an automatic side-effect.
        run(force_refresh_tles=refresh, render_doc=False)
